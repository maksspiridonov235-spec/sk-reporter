#!/usr/bin/env python3
"""Анализатор DOCX для отладки замены руководителя."""

from docx import Document
import sys

def analyze_document(filepath):
    """Показывает где сидит руководитель в документе."""
    doc = Document(filepath)
    
    print(f"Файл: {filepath}")
    print(f"Таблиц: {len(doc.tables)}\n")
    
    keywords = ['руковод', 'мандж', 'аниськ', 'и.о.', 'и.о', 'прораб', 'начальник']
    found_cells = []
    
    for t_idx, table in enumerate(doc.tables):
        table_has_match = False
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if any(k in text.lower() for k in keywords):
                    if not table_has_match:
                        print(f"\n=== ТАБЛИЦА {t_idx} ===")
                        table_has_match = True
                    
                    print(f"\n[Строка {r_idx}, Ячейка {c_idx}]:")
                    print(f"  Полный текст: {repr(text)}")
                    
                    # Детализация по параграфам и run'ам
                    print(f"  Структура:")
                    for p_idx, para in enumerate(cell.paragraphs):
                        if para.text.strip():
                            print(f"    Параграф {p_idx}: {repr(para.text)}")
                            runs = [repr(r.text) for r in para.runs if r.text]
                            if runs:
                                print(f"      Runs: {runs}")
                    
                    found_cells.append({
                        'table': t_idx,
                        'row': r_idx,
                        'cell': c_idx,
                        'text': text,
                        'cell_obj': cell
                    })
    
    return found_cells

def main():
    filepath = "Ежедневный отчет (ЮНС) от 26.04.2026 г. (БКНС-4) Пряхин И.Н..docx"
    found = analyze_document(filepath)
    
    print(f"\n\n=== ИТОГО НАЙДЕНО: {len(found)} ячеек ===")
    for cell in found:
        print(f"  [{cell['table']},{cell['row']},{cell['cell']}]: {cell['text'][:50]}...")

if __name__ == "__main__":
    main()
