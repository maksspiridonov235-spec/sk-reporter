"""
Логика обработки .docx файлов.
Все функции — аналоги VBA-макросов из Word, переписанные на python-docx.
"""

import os
import re
from pathlib import Path
import shutil
import zipfile
import xml.etree.ElementTree as ET
from copy import deepcopy
from datetime import datetime, timedelta
from typing import Literal, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from lxml import etree

from sk_reporter.companies import COMPANIES
from sk_reporter.template_layout import apply_layout

# Regex для поиска дат в разных форматах
_DATE_RE = re.compile(
    r"\d{1,2}\s*[./\-]\s*\d{1,2}\s*[./\-]\s*\d{4}\s*г\.?|"  # DD.MM.YYYY с буквой г
    r"\d{1,2}[./\-]\d{1,2}[./\-]\d{4}\s*г\.?|"               # DD.MM.YYYY с буквой г
    r"\d{1,2}\s*[./\-]\s*\d{1,2}\s*[./\-]\s*\d{4}|"          # DD.MM.YYYY
    r"\d{1,2}[./\-]\d{1,2}[./\-]\d{4}|"                      # DD.MM.YYYY
    r"\d{1,2}\s*[./\-]\s*\d{1,2}\s*[./\-]\s*\d{2}\s*г\.?|"   # DD.MM.YY с буквой г
    r"\d{1,2}[./\-]\d{1,2}[./\-]\d{2}\s*г\.?|"               # DD.MM.YY с буквой г
    r"\d{1,2}\s*[./\-]\s*\d{1,2}\s*[./\-]\s*\d{2}|"          # DD.MM.YY
    r"\d{1,2}[./\-]\d{1,2}[./\-]\d{2}"                       # DD.MM.YY
)

# ── Макрос 1: HighlightSecondRow_No5991 ────────────────────────────────────

def highlight_second_row(doc: Document) -> int:
    """
    Для каждой таблицы документа: если строк >= 2,
    красит 2-ю строку в голубой (#BDD6EE), жирный, по центру.
    Работает с объединёнными ячейками через XML напрямую.
    """
    BLUE_HEX = "BDD6EE"
    processed = 0

    for tbl in doc.tables:
        if len(tbl.rows) < 2:
            continue

        # Берём реальные <w:tc> элементы второй строки из XML (без дублей от merge)
        row_el = tbl.rows[1]._tr
        seen_tc = set()
        for tc in row_el.findall(qn("w:tc")):
            if id(tc) in seen_tc:
                continue
            seen_tc.add(id(tc))

            # Пропускаем ячейки с vMerge (продолжение вертикального объединения)
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                vMerge = tcPr.find(qn("w:vMerge"))
                if vMerge is not None and vMerge.get(qn("w:val")) != "restart":
                    continue

            # Фон
            if tcPr is None:
                tcPr = etree.SubElement(tc, qn("w:tcPr"))
            shd = tcPr.find(qn("w:shd"))
            if shd is None:
                shd = etree.SubElement(tcPr, qn("w:shd"))
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), BLUE_HEX)
            # Убираем тему чтобы наш цвет не перекрывался
            for attr in list(shd.attrib):
                if "theme" in attr.lower():
                    del shd.attrib[attr]

            # Вертикальное выравнивание
            vAlign = tcPr.find(qn("w:vAlign"))
            if vAlign is None:
                vAlign = etree.SubElement(tcPr, qn("w:vAlign"))
            vAlign.set(qn("w:val"), "center")

            # Текст: жирный, по центру
            for p in tc.findall(qn("w:p")):
                pPr = p.find(qn("w:pPr"))
                if pPr is None:
                    pPr = etree.SubElement(p, qn("w:pPr"))
                jc = pPr.find(qn("w:jc"))
                if jc is None:
                    jc = etree.SubElement(pPr, qn("w:jc"))
                jc.set(qn("w:val"), "center")
                for r in p.findall(qn("w:r")):
                    rPr = r.find(qn("w:rPr"))
                    if rPr is None:
                        rPr = etree.SubElement(r, qn("w:rPr"))
                    b = rPr.find(qn("w:b"))
                    if b is None:
                        etree.SubElement(rPr, qn("w:b"))

        processed += 1

    return processed


# ── Картинки ────────────────────────────────────────────────────────────────

IMG_WIDTH_CM = 5.33
IMG_HEIGHT_CM = 4.0
EMU_PER_CM = 360000
DRAWING_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
EMU_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"

PREPARE_PIPELINE_ID = "2026-06-02-layout-full"


def _iter_all_story_xml_roots(doc: Document) -> list:
    """Корни основного текста и колонтитулов (без повторной обработки linked)."""
    roots: list = []
    seen: set[int] = set()

    def add(el) -> None:
        if el is None:
            return
        eid = id(el)
        if eid in seen:
            return
        seen.add(eid)
        roots.append(el)

    add(doc.element.body)
    for section in doc.sections:
        for attr in (
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ):
            part = getattr(section, attr, None)
            if part is None:
                continue
            if getattr(part, "is_linked_to_previous", False):
                continue
            add(getattr(part, "_element", None))
    return roots


def resize_inline_images(doc: Document) -> int:
    """Все инлайн-картинки → 5,33 × 4 см (как старый NewMacros)."""
    img_w = int(IMG_WIDTH_CM * EMU_PER_CM)
    img_h = int(IMG_HEIGHT_CM * EMU_PER_CM)
    count = 0
    for shape in doc.inline_shapes:
        count += 1
        drawing = shape._inline
        extent = drawing.find(f"{{{DRAWING_NS}}}extent")
        if extent is not None:
            extent.set("cx", str(img_w))
            extent.set("cy", str(img_h))
        for ext in drawing.iter(f"{{{EMU_NS}}}ext"):
            ext.set("cx", str(img_w))
            ext.set("cy", str(img_h))
    return count


# ── Шрифт и макеты (без ширины таблиц, картинок и заливки) ─────────────────

FONT_NAME = "Times New Roman"
FONT_SIZE_HALF_POINTS = "20"  # 10 pt
MIN_ROW_HEIGHT_CM = 0.6


def _iter_document_paragraphs(doc: Document):
    """Все абзацы документа (тело, таблицы, колонтитулы) через XML w:p."""
    from docx.text.paragraph import Paragraph

    parent = doc.element.body
    for root in _iter_all_story_xml_roots(doc):
        for p_el in root.iter(qn("w:p")):
            yield Paragraph(p_el, parent)


def _apply_font_to_paragraph(para) -> None:
    for run in para.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(10)
    pPr = para._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = etree.SubElement(pPr, qn("w:rPr"))
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = etree.SubElement(rPr, qn("w:sz"))
    sz.set(qn("w:val"), FONT_SIZE_HALF_POINTS)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = etree.SubElement(rPr, qn("w:szCs"))
    szCs.set(qn("w:val"), FONT_SIZE_HALF_POINTS)


def format_fonts_only(doc: Document) -> None:
    """Только шрифт Times New Roman 10 pt."""
    for para in _iter_document_paragraphs(doc):
        _apply_font_to_paragraph(para)


def _reset_paragraph_layout_element(p_el) -> None:
    """Обнуление отступов и интервалов одного w:p (прямое форматирование)."""
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = etree.SubElement(p_el, qn("w:pPr"))
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = etree.SubElement(pPr, qn("w:spacing"))
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:beforeAutospacing"), "0")
    spacing.set(qn("w:afterAutospacing"), "0")
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = etree.SubElement(pPr, qn("w:ind"))
    for key in (
        "left",
        "right",
        "firstLine",
        "hanging",
        "start",
        "end",
        "firstLineChars",
        "hangingChars",
    ):
        ind.set(qn(f"w:{key}"), "0")


def reset_paragraph_layout(doc: Document) -> int:
    """Обнуление отступов и интервалов всех абзацев (тело + колонтитулы)."""
    count = 0
    for root in _iter_all_story_xml_roots(doc):
        for p_el in root.iter(qn("w:p")):
            _reset_paragraph_layout_element(p_el)
            count += 1
    return count


def apply_table_geometry(doc: Document, min_height_cm: float = MIN_ROW_HEIGHT_CM) -> None:
    """Минимальная высота строк и вертикальное выравнивание ячеек по центру."""
    row_h = str(int(min_height_cm * 567))
    for table in doc.tables:
        for row in table.rows:
            tr = row._tr
            tr_pr = tr.find(qn("w:trPr"))
            if tr_pr is None:
                tr_pr = etree.SubElement(tr, qn("w:trPr"))
            tr_h = tr_pr.find(qn("w:trHeight"))
            if tr_h is None:
                tr_h = etree.SubElement(tr_pr, qn("w:trHeight"))
            tr_h.set(qn("w:val"), row_h)
            tr_h.set(qn("w:hRule"), "atLeast")
            for tc in tr.findall(qn("w:tc")):
                tc_pr = tc.find(qn("w:tcPr"))
                if tc_pr is None:
                    tc_pr = etree.SubElement(tc, qn("w:tcPr"))
                    tc.insert(0, tc_pr)
                v_align = tc_pr.find(qn("w:vAlign"))
                if v_align is None:
                    v_align = etree.SubElement(tc_pr, qn("w:vAlign"))
                v_align.set(qn("w:val"), "center")


# ── Макросы 3 и 4: ReplaceDateInReportLine / ReplaceDateInReportLine2 ──────

def replace_date_in_report_line(doc: Document, mode: Literal["today", "yesterday"] = "today", target_date: Optional[str] = None) -> bool:
    """
    Заменяет дату в ячейке [2,1] таблицы на сегодняшнюю или вчерашнюю.
    Работает для всех типов отчётов (геодезия, обычные, несгруппированные).
    Возвращает True если замена выполнена.
    """
    if target_date is None:
        target_date = (
            datetime.now().strftime("%d.%m.%Y")
            if mode == "today"
            else (datetime.now() - timedelta(days=1)).strftime("%d.%m.%Y")
        )

    # Заменяем дату в ячейке [2,1] первой таблицы
    if len(doc.tables) > 0:
        table = doc.tables[0]
        if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
            cell = table.rows[2].cells[1]
            for para in cell.paragraphs:
                # Собираем весь текст параграфа (дата может быть разбита на много runs)
                full_text = "".join(run.text for run in para.runs)
                if _DATE_RE.search(full_text):
                    # Заменяем дату, очищаем все runs и пишем результат в первый run
                    new_text = _DATE_RE.sub(target_date, full_text)
                    if para.runs:
                        para.runs[0].text = new_text
                        # Удаляем остальные runs
                        for run in para.runs[1:]:
                            r = run._element
                            r.getparent().remove(r)
                    return True

    return False


def prepare_template_with_date(template_path: str, work_dir: str | None = None) -> str:
    """
    Копирует шаблон и меняет в нём дату на сегодняшнюю.
    Возвращает путь к обновленному временному шаблону.
    """
    import tempfile

    if work_dir is None:
        work_dir = tempfile.gettempdir()

    # Копируем шаблон во временный файл
    temp_template = os.path.join(work_dir, f"template_{os.path.basename(template_path)}")
    shutil.copy2(template_path, temp_template)

    # Меняем дату в копии
    doc = Document(temp_template)
    replace_date_in_report_line(doc, "today")
    doc.save(temp_template)

    return temp_template


# ── Объединение отчётов ─────────────────────────────────────────────────────

def _zip_replace(zip_path: str, inner_name: str, new_data: bytes) -> None:
    """Заменяет файл внутри ZIP-архива."""
    tmp = zip_path + ".tmp"
    with zipfile.ZipFile(zip_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == inner_name:
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
    os.replace(tmp, zip_path)


_EXT_TO_CONTENT_TYPE = {
    ".jpeg": "image/jpeg",
    ".jpg":  "image/jpeg",
    ".png":  "image/png",
    ".gif":  "image/gif",
    ".bmp":  "image/bmp",
    ".tiff": "image/tiff",
    ".wmf":  "image/x-wmf",
    ".emf":  "image/x-emf",
}


def _copy_media_from_docx(src_path: str, dst_path: str, remap: dict) -> None:
    """
    Копирует медиафайлы из src в dst, обновляет [Content_Types].xml в dst.
    Заполняет remap {старое_имя: новое_имя}.
    """
    with zipfile.ZipFile(src_path, "r") as src_zip:
        src_media = [n for n in src_zip.namelist() if n.startswith("word/media/")]
        if not src_media:
            return

        # Читаем текущий [Content_Types].xml из dst
        with zipfile.ZipFile(dst_path, "r") as dst_zip:
            existing = set(dst_zip.namelist())
            ct_xml = dst_zip.read("[Content_Types].xml").decode("utf-8")

        ct_root = ET.fromstring(ct_xml)
        ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
        ET.register_namespace("", ct_ns)

        # Собираем уже объявленные расширения
        declared_exts = {
            el.get("Extension", "").lower()
            for el in ct_root.findall(f"{{{ct_ns}}}Default")
        }

        new_media: list[tuple[str, bytes]] = []  # (новое_имя_в_zip, данные)

        for src_name in src_media:
            base = os.path.basename(src_name)
            stem, ext = os.path.splitext(base)
            candidate = f"word/media/{base}"
            counter = 1
            while candidate in existing:
                candidate = f"word/media/{stem}_{counter}{ext}"
                counter += 1
            data = src_zip.read(src_name)
            new_media.append((candidate, data))
            existing.add(candidate)
            remap[base] = os.path.basename(candidate)

            # Добавляем Default в [Content_Types].xml если расширение новое
            ext_lower = ext.lower().lstrip(".")
            if ext_lower not in declared_exts:
                content_type = _EXT_TO_CONTENT_TYPE.get(ext.lower(), "application/octet-stream")
                new_default = ET.SubElement(ct_root, f"{{{ct_ns}}}Default")
                new_default.set("Extension", ext_lower)
                new_default.set("ContentType", content_type)
                declared_exts.add(ext_lower)

        # Записываем медиафайлы и обновлённый [Content_Types].xml в dst
        new_ct_xml = ET.tostring(ct_root, encoding="unicode", xml_declaration=False)
        new_ct_xml = "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + new_ct_xml

        with zipfile.ZipFile(dst_path, "a") as dst_zip:
            for name, data in new_media:
                dst_zip.writestr(name, data)

        _zip_replace(dst_path, "[Content_Types].xml", new_ct_xml.encode("utf-8"))


def _fix_image_refs_in_element(element, remap: dict) -> None:
    """
    Обновляет все ссылки на картинки в XML-элементе согласно remap.
    Картинки в docx ссылаются через r:embed в <a:blip> и <v:imagedata>.
    Remap применяется к именам файлов в relationships, но мы патчим
    уже resolved-имена прямо в XML через rId → имя файла.
    Этот патч применяется на уровне имён файлов в rels.
    """
    pass  # Rels патчим отдельно в merge_reports


def merge_reports(template_path: str, report_paths: list[str], output_path: str) -> int:
    """
    Объединяет отчёты из report_paths в шаблон template_path,
    сохраняет результат в output_path.
    Корректно переносит картинки через ZIP чтобы они не были битыми.
    Возвращает количество вставленных файлов.
    """
    shutil.copy2(template_path, output_path)
    inserted = 0

    for path in sorted(report_paths):
        try:
            src_doc = Document(path)
        except Exception:
            continue

        # 1. Копируем медиафайлы из src в output, получаем карту переименований
        remap: dict[str, str] = {}
        _copy_media_from_docx(path, output_path, remap)

        # 2. Читаем relationships из источника чтобы построить rId → имя файла
        src_rels: dict[str, str] = {}
        with zipfile.ZipFile(path, "r") as zf:
            rels_name = "word/_rels/document.xml.rels"
            if rels_name in zf.namelist():
                rels_xml = zf.read(rels_name).decode("utf-8")
                root = ET.fromstring(rels_xml)
                ns = "http://schemas.openxmlformats.org/package/2006/relationships"
                for rel in root.findall(f"{{{ns}}}Relationship"):
                    rid = rel.get("Id", "")
                    target = rel.get("Target", "")
                    # target вида "../media/image1.png" или "media/image1.png"
                    basename = os.path.basename(target)
                    src_rels[rid] = basename

        # 3. Добавляем relationships в output для новых медиафайлов
        #    и строим карту старых rId → новых rId
        rid_remap: dict[str, str] = {}
        if remap:
            with zipfile.ZipFile(output_path, "r") as zf:
                rels_name = "word/_rels/document.xml.rels"
                ET.register_namespace("", "http://schemas.openxmlformats.org/package/2006/relationships")
                rels_xml = zf.read(rels_name).decode("utf-8")
                rels_root = ET.fromstring(rels_xml)

            ns = "http://schemas.openxmlformats.org/package/2006/relationships"
            img_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"

            # Найдём максимальный rId в output
            existing_ids = set()
            for rel in rels_root.findall(f"{{{ns}}}Relationship"):
                existing_ids.add(rel.get("Id", ""))
            rid_counter = 1
            while f"rId{rid_counter}" in existing_ids:
                rid_counter += 1

            # Для каждого rId из источника с картинкой — добавляем новый rel в output
            for old_rid, old_basename in src_rels.items():
                if old_basename not in remap:
                    continue
                new_basename = remap[old_basename]
                new_rid = f"rId{rid_counter}"
                while new_rid in existing_ids:
                    rid_counter += 1
                    new_rid = f"rId{rid_counter}"
                rid_remap[old_rid] = new_rid
                existing_ids.add(new_rid)
                rid_counter += 1
                new_rel = ET.SubElement(rels_root, f"{{{ns}}}Relationship")
                new_rel.set("Id", new_rid)
                new_rel.set("Type", img_type)
                new_rel.set("Target", f"media/{new_basename}")

            # Записываем обновлённый rels обратно в ZIP
            new_rels_xml = ET.tostring(rels_root, encoding="unicode", xml_declaration=False)
            new_rels_xml = '<?xml version=\'1.0\' encoding=\'UTF-8\' standalone=\'yes\'?>\n' + new_rels_xml
            _zip_replace(output_path, rels_name, new_rels_xml.encode("utf-8"))

        # 4. Копируем тело документа, патча rId в XML-элементах
        master = Document(output_path)

        # Собираем элементы тела, пропуская ведущие пустые параграфы с разрывами
        body_elements = [
            el for el in src_doc.element.body
            if (el.tag.split("}")[-1] if "}" in el.tag else el.tag) != "sectPr"
        ]

        # Отрезаем пустые page-break параграфы в начале
        def _is_page_break_para(el):
            tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            if tag != "p":
                return False
            from lxml import etree
            xml = etree.tostring(el, encoding="unicode")
            has_break = 'w:br' in xml and 'page' in xml
            # Текст — всё что не в тегах
            text = "".join(el.itertext()).strip()
            return has_break and not text

        while body_elements and _is_page_break_para(body_elements[0]):
            body_elements.pop(0)

        if inserted > 0 and body_elements:
            master.add_page_break()

        for element in body_elements:
            el_copy = deepcopy(element)
            if rid_remap:
                _patch_rids(el_copy, rid_remap)
            master.element.body.append(el_copy)

        master.save(output_path)
        inserted += 1

    return inserted


def _patch_rids(element, rid_remap: dict[str, str]) -> None:
    """Рекурсивно заменяет атрибуты r:embed, r:id, r:link на новые rId."""
    REMAP_ATTRS = (
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed",
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id",
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link",
    )
    for attr in REMAP_ATTRS:
        val = element.get(attr)
        if val and val in rid_remap:
            element.set(attr, rid_remap[val])
    for child in element:
        _patch_rids(child, rid_remap)


# ── Переименование файлов ───────────────────────────────────────────────────

def rename_files(folder: str, mode: Literal["today", "yesterday"]) -> list[str]:
    """
    Переименовывает .docx/.doc файлы в папке:
    заменяет все даты вида DD.MM.YYYY в имени файла на сегодняшнюю или вчерашнюю.
    Возвращает список строк лога.
    """
    new_date = (
        datetime.now().strftime("%d.%m.%Y")
        if mode == "today"
        else (datetime.now() - timedelta(days=1)).strftime("%d.%m.%Y")
    )
    log = []
    for filename in os.listdir(folder):
        if not any(filename.lower().endswith(ext) for ext in (".docx", ".doc")):
            continue
        new_name = _DATE_RE.sub(new_date, filename)
        if new_name != filename:
            try:
                os.rename(
                    os.path.join(folder, filename),
                    os.path.join(folder, new_name),
                )
                log.append(f"Переименован: {filename} → {new_name}")
            except Exception as e:
                log.append(f"Ошибка: {filename} — {e}")
    return log


# ── Переименование результатов ──────────────────────────────────────────────

def rename_results(folder: str, target_date: str) -> list[str]:
    """
    Переименовывает Евракор_merged.docx → Евракор_Ежедневный отчёт СК за DD.MM.YYYY.docx
    """
    log = []
    for filename in os.listdir(folder):
        if not filename.lower().endswith((".docx", ".doc")):
            continue
        if "_merged" not in filename:
            continue
        company = filename.replace("_merged.docx", "").replace("_merged.doc", "")
        ext = ".docx" if filename.lower().endswith(".docx") else ".doc"
        new_name = f"{company}_Ежедневный отчёт СК за {target_date}{ext}"
        filepath = os.path.join(folder, filename)
        try:
            doc = Document(filepath)
            replace_date_in_report_line(doc, target_date=target_date)
            doc.save(filepath)
            os.rename(filepath, os.path.join(folder, new_name))
            log.append(
                f"[OK] {new_name}: на диске {filename} → {new_name}; "
                f"в документе — дата {target_date}"
            )
        except Exception as e:
            log.append(f"[ERR] {filename}: {e}")
    return log


# Дата в шапке болванки: DD.MM.YYYYг
_BOLVANKA_DATE_RE = re.compile(r"\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4}\s*г")


def _write_paragraph_text(para, new_text: str) -> None:
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            r = run._element
            r.getparent().remove(r)
    else:
        para.add_run(new_text)


def _set_title_line_date(para, target_date: str) -> bool:
    """
    Строка шапки болванки «Отчёт … за DD.MM.YYYYг».
    Заменяет существующую дату или дописывает после «за», если даты ещё нет.
    """
    full_text = "".join(run.text for run in para.runs)
    stripped = full_text.strip()
    if not stripped or "за" not in stripped.lower():
        return False

    dated_suffix = f"{target_date}г"

    if _BOLVANKA_DATE_RE.search(full_text):
        new_text = _BOLVANKA_DATE_RE.sub(dated_suffix, full_text, count=1)
    elif stripped.endswith("за") or stripped.endswith("за "):
        new_text = full_text.rstrip()
        if not new_text.endswith(" "):
            new_text += " "
        new_text += dated_suffix
    else:
        m = re.search(r"(\bза)\s*\S*\s*$", full_text, re.IGNORECASE)
        if not m:
            return False
        new_text = full_text[: m.start(2)] + " " + dated_suffix

    _write_paragraph_text(para, new_text)
    return True


def rename_templates(folder: str, target_date: str) -> list[str]:
    """
    Ставит дату в болванках в тексте («Отчёт … за …» и таблица).
    Имя файла не меняется (болванки: «Компания.docx»).
    """
    log = []
    folder_path = Path(folder)

    for filepath in sorted(folder_path.iterdir()):
        if filepath.suffix.lower() not in (".docx", ".doc"):
            continue
        try:
            doc = Document(os.fspath(filepath))
            updated = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                if "отчёт" in text.lower() and "за" in text.lower():
                    if _set_title_line_date(para, target_date):
                        updated.append("заголовок")
                    break

            if replace_date_in_report_line(doc, target_date=target_date):
                updated.append("таблица")

            if not updated:
                log.append(
                    f"[ERR] {filepath.name}: в тексте нет строки «Отчёт … за» — дату не заменили"
                )
                continue

            doc.save(os.fspath(filepath))

            parts = ", ".join(updated)
            log.append(
                f"[OK] {filepath.name}: в тексте дата {target_date}г ({parts})"
            )
        except Exception as e:
            log.append(f"[ERR] {filepath.name}: {e}")
    return log


def prepare_report_file(filepath: str, layout: dict, target_date: str) -> tuple[bool, str]:
    try:
        doc = Document(filepath)
    except Exception as e:
        return False, str(e)
    parts = [f"версия {PREPARE_PIPELINE_ID}"]
    n = highlight_second_row(doc)
    parts.append(f"заливка {n}")
    format_fonts_only(doc)
    parts.append("шрифт")
    n_layout = reset_paragraph_layout(doc)
    parts.append(f"макеты {n_layout}")
    n_img = resize_inline_images(doc)
    parts.append(f"картинки {n_img}")
    if replace_date_in_report_line(doc, target_date=target_date):
        parts.append(f"дата в тексте {target_date}")
    else:
        parts.append("дата в тексте не найдена")
    apply_table_geometry(doc)
    parts.append("строки 0,6")
    layout_warns = apply_layout(doc, layout, only_main_table=True)
    parts.append("сетка")
    if layout_warns:
        parts.append("⚠ " + layout_warns[0])
        if len(layout_warns) > 1:
            parts.append(f"(+{len(layout_warns) - 1} предупр.)")
    doc.save(filepath)
    return True, ", ".join(parts)


def prepare_uploaded_reports(upload_dir: str, layout: dict, target_date: str) -> list[str]:
    log = []
    folder = Path(upload_dir)
    files = sorted(f for f in folder.iterdir() if f.suffix.lower() in (".docx", ".doc"))
    if not files:
        log.append("[ERR] Нет загруженных отчётов")
        return log
    for f in files:
        ok, msg = prepare_report_file(str(f), layout, target_date)
        log.append(f"[{'OK' if ok else 'ERR'}] {f.name}: {msg}")
    return log


def extract_report_data(filepath: str) -> dict:
    """Извлекает данные из отчёта для проверки описаний работ."""
    try:
        doc = Document(filepath)
        if not doc.tables:
            return None

        table = doc.tables[0]
        report_data = {
            "filename": Path(filepath).name,
            "works": []
        }

        # Парсим таблицу и ищем описания работ
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]

            # Проверяем на признаки строки с работой (начинается с номера)
            if cells and cells[0] and cells[0][0].isdigit() and "." in cells[0][:5]:
                work = {
                    "row": row_idx,
                    "description": cells[0] if cells else "",
                    "volumes": {}
                }

                # Ищем объёмы в этой строке
                full_text = "".join(cells)
                if "Проектный объем" in full_text:
                    work["volumes"]["project"] = _extract_volume(full_text, "Проектный объем")
                if "Объем за сутки" in full_text:
                    work["volumes"]["daily"] = _extract_volume(full_text, "Объем за сутки")
                if "Накопительный объем" in full_text:
                    work["volumes"]["cumulative"] = _extract_volume(full_text, "Накопительный объем")

                report_data["works"].append(work)

        return report_data
    except Exception as e:
        return {"error": str(e), "filename": Path(filepath).name}


def _extract_volume(text: str, label: str) -> str:
    """Извлекает значение объёма из текста."""
    pattern = rf"{label}\s*[–-]?\s*([^\.;,\n]+)"
    match = re.search(pattern, text)
    return match.group(1).strip() if match else ""
