"""AI Агент переключения руководителя через Ollama.
Использует LLM для анализа и замены.
"""

import ollama
from docx import Document
from pathlib import Path
from typing import Literal

MODEL = "gemma4:31b-cloud"

SYSTEM_PROMPT = """Ты - агент для редактирования отчётов строительного контроля.

Твоя задача: для КАЖДОЙ ячейки сказать — это шапка или подвал, и что туда вставить.

ПРАВИЛА:
1. Смотри на текст ячейки и её положение в документе
2. ШАПКА = подписи в верхней части документа (обычно рядом с таблицей, короткие)
3. ПОДВАЛ = подписи в нижней части документа (после таблицы, полные должности)
4. В шапке пишем КОРОТКИЙ заголовок: "Руководитель" или "И.О. Руководителя"
5. В подвале пишем ПОЛНУЮ должность: "Руководитель проекта СК" или "И.О. Руководителя проекта СК"
6. ФИО меняем везде одинаково

СТАРЫЙ РУКОВОДИТЕЛЬ (искать):
- ФИО: Аниськов Владимир Иванович, Манджиев Игорь Александрович (и опечатки)
- Должности: Руководитель, Руководитель проекта СК, И.о./И.О. Руководителя и т.д.

НОВЫЙ РУКОВОДИТЕЛЬ:
- Для Аниськова: ФИО "Аниськов Владимир Иванович", шапка "Руководитель", подвал "Руководитель проекта СК"
- Для Манджиева: ФИО "Манджиев Игорь Александрович", шапка "И.О. Руководителя", подвал "И.О. Руководителя проекта СК"

Ответь JSON — список ячеек с классификацией:
{
  "cells": [
    {"table": 0, "row": 8, "cell": 4, "section": "header", "old": "И.о. Руководителя", "new": "Руководитель"},
    {"table": 0, "row": 44, "cell": 4, "section": "footer", "old": "И.о. Руководителя проекта СК", "new": "Руководитель проекта СК"}
  ],
  "confidence": 0-100
}"""


def extract_text_from_docx(filepath: str) -> str:
    """Извлекает текст из DOCX для анализа."""
    try:
        doc = Document(filepath)
        parts = []

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and len(text) > 2:
                        parts.append(text)

        return "\n".join(parts[:100])
    except Exception as e:
        return f"Error: {e}"


def analyze_with_ai(filepath: str, to_leader: str) -> dict:
    """Старый метод для совместимости."""
    return analyze_with_ai_cells([], to_leader)


def analyze_with_ai_cells(cells_list: list, to_leader: str) -> dict:
    """AI классифицирует каждую ячейку — шапка или подвал, и что вставить."""
    
    if to_leader == "aniskov":
        target_fio = "Аниськов Владимир Иванович"
        target_title = "Руководитель"
        target_role = "Руководитель проекта СК"
        direction = "Манджиев → Аниськов (и.о. → руководитель)"
    else:
        target_fio = "Манджиев Игорь Александрович"
        target_title = "И.О. Руководителя"
        target_role = "И.О. Руководителя проекта СК"
        direction = "Аниськов → Манджиев (руководитель → и.о.)"

    # Формируем текст для AI — просто список ячеек
    cells_text = "\n".join([
        f"[T{c['table']} R{c['row']} C{c['cell']}]: {c['text']}" 
        for c in cells_list
    ])

    prompt = f"""Направление: {direction}

ЦЕЛЕВЫЕ ЗНАЧЕНИЯ:
- ФИО: {target_fio}
- ШАПКА (коротко): {target_title}
- ПОДВАЛ (полностью): {target_role}

Проанализируй КАЖДУЮ ячейку:
1. Это шапка (верх документа, подпись под таблицей) или подвал (низ, под таблицей)?
2. Что там написано про руководителя?
3. Что вставить (section определяет: header={target_title}, footer={target_role})

Ячейки:
{cells_text}

ВАЖНО: верни ТОЧНЫЕ old тексты из ячеек.

Ответь JSON:
{{
  "cells": [
    {{"table": 0, "row": 8, "cell": 4, "section": "header", "old": "И.о. Руководителя", "new": "{target_title}"}},
    {{"table": 0, "row": 44, "cell": 4, "section": "footer", "old": "И.о. Руководителя проекта СК", "new": "{target_role}"}}
  ],
  "confidence": 0-100
}}"""

    try:
        response = ollama.chat(
            model=MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": prompt}
            ],
            options={"temperature": 0.0, "num_predict": 2000},
        )

        answer = response["message"]["content"]

        import json
        import re

        json_match = re.search(r'\{.*\}', answer, re.DOTALL)
        if json_match:
            return json.loads(json_match.group())

        return {"error": "No JSON found", "raw": answer}

    except Exception as e:
        return {"error": str(e)}


def _replace_in_runs(cell, old_text: str, new_text: str) -> int:
    """Заменяет текст внутри run'ов ячейки, сохраняя форматирование."""
    changes = 0
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            # Проверяем точное совпадение после нормализации пробелов
            normalized_run = " ".join(run.text.split())  # убираем лишние пробелы
            if old_text in normalized_run or old_text in run.text:
                # Заменяем в оригинальном тексте
                run.text = run.text.replace(old_text, new_text)
                changes += 1
    return changes


def _switch_single_file(filepath: str, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """AI говорит что менять, мы меняем в cell.text целиком."""
    try:
        doc = Document(filepath)

        if not doc.tables:
            return False, "Нет таблиц в документе"

        # Собираем все ячейки с текстом и координатами
        cells_data = []
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        cells_data.append({
                            'table': t_idx, 'row': r_idx, 'cell': c_idx,
                            'text': text, 'cell_obj': cell
                        })

        # Отправляем AI список
        ai_result = analyze_with_ai_cells(cells_data, leader)
        
        if "error" in ai_result:
            return False, f"AI ошибка: {ai_result['error']}"

        cell_replacements = ai_result.get("cells", [])
        
        if not cell_replacements:
            return False, "AI не нашёл ячеек для замены"

        changes = 0
        
        # Применяем замены по координатам ячеек
        for repl in cell_replacements:
            t = repl.get('table')
            r = repl.get('row')
            c_idx = repl.get('cell')
            old_text = repl.get('old', '')
            new_text = repl.get('new', '')
            
            if not old_text or not new_text:
                continue
            
            # Находим ячейку по координатам
            for cell_info in cells_data:
                if (cell_info['table'] == t and 
                    cell_info['row'] == r and 
                    cell_info['cell'] == c_idx):
                    
                    # Заменяем весь текст ячейки
                    if old_text in cell_info['text']:
                        cell_info['cell_obj'].text = cell_info['text'].replace(old_text, new_text)
                        changes += 1
                    break

        if changes == 0:
            return False, f"AI дал {len(cell_replacements)} ячеек, но не применилось"

        doc.save(filepath)
        return True, f"→ {Path(filepath).name}: замен {changes}"

    except Exception as e:
        return False, f"→ {Path(filepath).name}: ошибка - {str(e)}"


def switch_leader(filepath: str, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """Переключает руководителя в одном файле (совместимость со старым API)."""
    return _switch_single_file(filepath, leader)


def switch_leader_ai(filepaths: list, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """Обрабатывает список файлов."""
    if not filepaths:
        return False, "Нет файлов для обработки"
    
    results = []
    success_count = 0
    total_changes = 0
    
    for filepath in filepaths:
        ok, msg = _switch_single_file(filepath, leader)
        results.append(msg)
        if ok:
            success_count += 1
            try:
                if "замен " in msg:
                    changes_str = msg.split("замен ")[-1].strip()
                    total_changes += int(changes_str)
            except:
                pass
    
    if success_count == 0:
        return False, "Ни один файл не обработан: " + "; ".join(results)
    
    output = "\n".join(results)
    output += f"\nОбработано: {success_count}/{len(filepaths)} файлов, замен: {total_changes}"
    return True, output
