"""
Агент инъекции: берёт оригинальный docx + исправленный текст от check_agent
и вставляет две строки под заголовками секции СК
(«Описание действий», «Участок, ПК», «Ссылка»):
первая — ЧАСТЬ 1 + 3 + 4, вторая — ЧАСТЬ 2.
Текст инженера в существующих строках не трогается.
"""

import re
import shutil
import tempfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Row
from lxml import etree

MODEL = "gemma4:31b-cloud"

FIXED_DOWNLOAD_SUFFIX = "_исправлен.docx"


def _split_part(search_text: str, num: int, next_num: int | None) -> list[str]:
    if next_num is not None:
        pattern = rf"ЧАСТЬ\s*{num}[^:\n]*[:\n](.*?)(?=ЧАСТЬ\s*{next_num}\b)"
    else:
        pattern = rf"ЧАСТЬ\s*{num}[^:\n]*[:\n](.*?)$"
    m = re.search(pattern, search_text, re.DOTALL | re.IGNORECASE)
    if not m:
        return []
    return [l.rstrip() for l in m.group(1).strip().splitlines()]


def _parse_parts(corrected_text: str):
    """Parse parts 1–4 from LLM output."""
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", corrected_text)

    section_match = re.search(
        r"##\s*ИСПРАВЛЕННЫЙ\s*ОТЧЁТ[^\n]*\n(.*?)$", cleaned, re.DOTALL | re.IGNORECASE
    )
    search_text = section_match.group(1).strip() if section_match else cleaned.strip()

    part1_lines = _split_part(search_text, 1, 2)
    part2_lines = _split_part(search_text, 2, 3)
    part3_lines = _split_part(search_text, 3, 4)
    part4_lines = _split_part(search_text, 4, None)

    if not part1_lines and not part2_lines:
        p2_start = re.search(
            r"(Наряд.допуск|Работы ведутся)", search_text, re.IGNORECASE
        )
        if p2_start:
            raw_part1 = search_text[: p2_start.start()].strip()
            raw_part2 = search_text[p2_start.start() :].strip()
            raw_part1 = re.sub(r"^ЧАСТЬ\s*1[^\n]*\n", "", raw_part1, flags=re.IGNORECASE).strip()
            part1_lines = [l.rstrip() for l in raw_part1.splitlines()] if raw_part1 else []
            part2_lines = [l.rstrip() for l in raw_part2.splitlines()] if raw_part2 else []

    print(
        f"[INJECT_AGENT] parsed part1={len(part1_lines)} part2={len(part2_lines)} "
        f"part3={len(part3_lines)} part4={len(part4_lines)} lines"
    )
    return part1_lines, part2_lines, part3_lines, part4_lines


def _cell_header_label(cell) -> str:
    text = cell.text.strip()
    return text.split("\n")[0].strip() if text else ""


def _unique_row_cells(row):
    seen = set()
    cells = []
    for cell in row.cells:
        tid = id(cell._tc)
        if tid in seen:
            continue
        seen.add(tid)
        cells.append(cell)
    return cells


def _table_trs(table) -> list:
    return table._tbl.findall(qn("w:tr"))


def _row_at(table, tr_index: int) -> _Row:
    return _Row(_table_trs(table)[tr_index], table)


def _column_role_indices(header_row) -> dict[str, int]:
    unique = _unique_row_cells(header_row)
    roles: dict[str, int] = {"description": 0}
    for i, cell in enumerate(unique):
        label = _cell_header_label(cell)
        if label.startswith("Участок"):
            roles["location"] = i
        elif label == "Ссылка":
            roles["reference"] = i
    return roles


def _find_sk_section_header_row(doc: Document):
    """Таблица, индекс <w:tr> заголовков и индексы колонок description/location/reference."""
    for table in doc.tables:
        trs = _table_trs(table)
        for ri, tr in enumerate(trs):
            row = _Row(tr, table)
            if row.cells[0].text.strip() != "Описание действий":
                continue
            if ri > 0:
                prev = _Row(trs[ri - 1], table).cells[0].text.strip().upper()
                if "СТРОИТЕЛЬНОГО КОНТРОЛЯ" not in prev:
                    continue
            return table, ri, _column_role_indices(row)
    return None, None, None


def _find_sk_section_header_cells(doc: Document):
    """Ячейки-заголовки строки секции СК (для report_builder и совместимости)."""
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            if row.cells[0].text.strip() != "Описание действий":
                continue
            if ri > 0:
                prev = table.rows[ri - 1].cells[0].text.strip().upper()
                if "СТРОИТЕЛЬНОГО КОНТРОЛЯ" not in prev:
                    continue
            cells = {"description": row.cells[0]}
            seen = {id(row.cells[0]._tc)}
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                label = _cell_header_label(cell)
                if label.startswith("Участок"):
                    cells["location"] = cell
                elif label == "Ссылка":
                    cells["reference"] = cell
            return cells, (ti, ri)
    return None, None


def _tr_index(table, tr) -> int:
    return _table_trs(table).index(tr)


def _insert_row_after(table, row_index: int, template_row_index: int) -> _Row:
    """Вставляет копию template_row сразу после row_index (через addnext, не tbl.insert)."""
    trs = _table_trs(table)
    new_tr = deepcopy(trs[template_row_index])
    trs[row_index].addnext(new_tr)
    return _Row(new_tr, table)


def _cells_for_roles(row, role_indices: dict[str, int]) -> dict:
    unique = _unique_row_cells(row)
    cells = {"description": unique[role_indices["description"]]}
    loc_i = role_indices.get("location")
    if loc_i is not None and loc_i < len(unique):
        cells["location"] = unique[loc_i]
    ref_i = role_indices.get("reference")
    if ref_i is not None and ref_i < len(unique):
        cells["reference"] = unique[ref_i]
    return cells


def _set_cell_vertical_align_top(cell) -> None:
    tc = cell._tc
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tc_pr)
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is None:
        v_align = etree.SubElement(tc_pr, qn("w:vAlign"))
    v_align.set(qn("w:val"), "top")


def _set_row_vertical_align_top(row) -> None:
    seen: set[int] = set()
    for cell in row.cells:
        tid = id(cell._tc)
        if tid in seen:
            continue
        seen.add(tid)
        _set_cell_vertical_align_top(cell)


def _write_lines_to_cell(cell, lines: list):
    """Заменяет содержимое ячейки, сохраняя первый параграф-заголовок."""
    if not lines:
        return
    tc = cell._tc
    paras = tc.findall(qn("w:p"))
    for p in paras[1:]:
        tc.remove(p)
    for line in lines:
        if line.strip():
            cell.add_paragraph(line)


def _write_lines_to_cell_data(cell, lines: list):
    """Полностью заменяет содержимое ячейки данными (без сохранения заголовка)."""
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    for line in lines:
        if line.strip():
            cell.add_paragraph(line)


def inject_into_docx(filepath: str, corrected_text: str, source_filename: str) -> dict:
    stem = Path(source_filename).stem
    try:
        print(f"[INJECT_AGENT] === FULL corrected_text ===\n{corrected_text[:500]}\n... (truncated) ===")
        part1_lines, part2_lines, part3_lines, part4_lines = _parse_parts(corrected_text)

        if not part1_lines and not part2_lines:
            return {
                "ok": False,
                "error": "Не удалось распарсить ЧАСТЬ 1 / ЧАСТЬ 2 из ответа агента",
                "docx_path": None,
            }

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir) / Path(filepath).name
            shutil.copy2(filepath, tmp_path)

            doc = Document(str(tmp_path))
            table, header_ri, role_indices = _find_sk_section_header_row(doc)
            if table is None or header_ri is None:
                return {
                    "ok": False,
                    "error": "Не найдена строка заголовков секции СК в документе",
                    "docx_path": None,
                }

            trs = _table_trs(table)
            template_ri = header_ri + 1 if header_ri + 1 < len(trs) else header_ri
            header_label_before = _row_at(table, header_ri).cells[0].text.strip()
            row_part1 = _insert_row_after(table, header_ri, template_ri)
            _set_row_vertical_align_top(row_part1)
            header_label_after = _row_at(table, header_ri).cells[0].text.strip()
            if header_label_before != header_label_after:
                return {
                    "ok": False,
                    "error": "Строка заголовков изменилась при вставке — операция отменена",
                    "docx_path": None,
                }
            print(
                f"[INJECT_AGENT] Inserted row 1 (parts 1,3,4) after header tr[{header_ri}], "
                f"template tr[{template_ri}], header preserved={header_label_after!r}"
            )

            cells1 = _cells_for_roles(row_part1, role_indices)
            _write_lines_to_cell_data(cells1["description"], part1_lines)
            print("[INJECT_AGENT] Wrote part 1 to row 1, 'Описание действий' column")

            if part3_lines and "location" in cells1:
                _write_lines_to_cell_data(cells1["location"], part3_lines)
                print("[INJECT_AGENT] Wrote part 3 to row 1, 'Участок, ПК' column")
            if part4_lines and "reference" in cells1:
                _write_lines_to_cell_data(cells1["reference"], part4_lines)
                print("[INJECT_AGENT] Wrote part 4 to row 1, 'Ссылка' column")

            if part2_lines:
                row1_ri = _tr_index(table, row_part1._tr)
                template_after_row1 = template_ri + 1
                row_part2 = _insert_row_after(table, row1_ri, template_after_row1)
                _set_row_vertical_align_top(row_part2)
                cells2 = _cells_for_roles(row_part2, role_indices)
                _write_lines_to_cell_data(cells2["description"], part2_lines)
                if "location" in cells2:
                    _write_lines_to_cell_data(cells2["location"], [])
                if "reference" in cells2:
                    _write_lines_to_cell_data(cells2["reference"], [])
                print(
                    f"[INJECT_AGENT] Inserted row 2 (part 2) after tr[{row1_ri}], "
                    "wrote to 'Описание действий' column"
                )

            dest = Path(filepath).resolve()
            doc.save(str(dest))

        print(f"[INJECT_AGENT] Saved to upload: {dest}")
        return {
            "ok": True,
            "docx_path": str(dest),
            "download_name": f"{stem}{FIXED_DOWNLOAD_SUFFIX}",
        }

    except Exception as e:
        import traceback

        print(f"[INJECT_AGENT] error: {e}\n{traceback.format_exc()}")
        return {"ok": False, "error": str(e), "docx_path": None}
