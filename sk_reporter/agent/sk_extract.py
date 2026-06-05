"""Извлечение секции СК и работ из docx."""

import re
from docx import Document

from sk_reporter.agent.report_parts import (
    is_zero_daily_volume,
    parse_numbered_items,
    parse_works_from_part1_lines,
    split_part1_part2_text,
)


def _cell_header_label(cell) -> str:
    text = cell.text.strip()
    return text.split("\n")[0].strip() if text else ""


def _cell_body_text(cell) -> str:
    text = cell.text.strip()
    lines = text.split("\n")
    label = _cell_header_label(cell)
    if lines and label and lines[0].strip() == label:
        body = "\n".join(lines[1:]).strip()
        if body:
            return body
    if label and text == label:
        return ""
    if label and text.startswith(label):
        rest = text[len(label) :].strip()
        return rest
    return text


def _unique_row_cells(row) -> list[str]:
    seen = set()
    cols = []
    for cell in row.cells:
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        cols.append(cell.text.strip())
    return cols


def find_sk_section_header_cells(doc: Document):
    """Ячейки-заголовки строки секции СК: описание, участок, ссылка."""
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            if row.cells[0].text.strip() != "Описание действий":
                continue
            if ri > 0:
                prev = table.rows[ri - 1].cells[0].text.strip().upper()
                if "СТРОИТЕЛЬНОГО КОНТРОЛЯ" not in prev:
                    continue
            cells = {"description": row.cells[0]}
            seen = {id(row.cells[0]._tc)}
            for cell in row.cells:
                if id(cell._tc) in seen:
                    continue
                seen.add(id(cell._tc))
                label = _cell_header_label(cell)
                if label.startswith("Участок"):
                    cells["location"] = cell
                elif label == "Ссылка":
                    cells["reference"] = cell
            return cells, (ti, ri)
    return None, None


def extract_full_text(filepath: str) -> str:
    """Весь текст docx: параграфы и строки таблиц."""
    try:
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        print(f"[SK_EXTRACT] extract_full_text error: {e}")
        return ""


def extract_sk_section(filepath: str) -> str:
    """Строки секции СК с колонками Описание / Участок, ПК / Ссылка."""
    try:
        doc = Document(filepath)
        lines = []
        for table in doc.tables:
            in_section = False
            for ri, row in enumerate(table.rows):
                c0 = row.cells[0].text.strip()
                if c0 == "Описание действий" and ri > 0:
                    prev = table.rows[ri - 1].cells[0].text.strip().upper()
                    if "СТРОИТЕЛЬНОГО КОНТРОЛЯ" in prev:
                        in_section = True
                        lines.append("--- Секция строительного контроля (таблица) ---")
                if not in_section:
                    continue
                if ri > 0 and c0.upper().startswith("РЕЗУЛЬТАТ ДУБЛИРУЮЩЕГО"):
                    break
                cols = _unique_row_cells(row)
                if any(c.strip() for c in cols):
                    lines.append(" | ".join(c.replace("\n", " | ") for c in cols))
        return "\n".join(lines)
    except Exception as e:
        print(f"[SK_EXTRACT] extract_sk_section error: {e}")
        return ""


def _end_before_sk_table(full_text: str, start: int) -> int:
    end = len(full_text)
    for pat in (
        r"РЕЗУЛЬТАТ\s+СТРОИТЕЛЬНОГО\s+КОНТРОЛЯ",
        r"Описание действий\s*\|",
        r"РЕЗУЛЬТАТ\s+ДУБЛИРУЮЩЕГО",
    ):
        em = re.search(pat, full_text[start:], re.IGNORECASE)
        if em:
            end = min(end, start + em.start())
    return end


def _extract_report_body_from_full_text(full_text: str) -> str:
    """Блок с работами и объёмами до таблицы СК."""
    m = re.search(r"Инспекционный контроль", full_text, re.IGNORECASE)
    if m:
        start = m.start()
        return full_text[start : _end_before_sk_table(full_text, start)].strip()

    m_end = re.search(r"РЕЗУЛЬТАТ\s+СТРОИТЕЛЬНОГО\s+КОНТРОЛЯ", full_text, re.IGNORECASE)
    if m_end and re.search(r"Проектный\s+объ[её]м", full_text[: m_end.start()], re.IGNORECASE):
        return full_text[: m_end.start()].strip()
    return ""


def _extract_sk_table_rows(doc: Document, coords: tuple[int, int]) -> list[dict]:
    """Строки таблицы под заголовками СК (описание / участок / ссылка по колонкам)."""
    ti, header_ri = coords
    table = doc.tables[ti]
    rows: list[dict] = []
    for ri in range(header_ri + 1, len(table.rows)):
        row = table.rows[ri]
        c0 = row.cells[0].text.strip()
        if c0.upper().startswith("РЕЗУЛЬТАТ ДУБЛИРУЮЩЕГО"):
            break
        cols = _unique_row_cells(row)
        if not cols:
            continue
        desc = cols[0] if len(cols) > 0 else ""
        loc = cols[1] if len(cols) > 1 else ""
        ref = cols[2] if len(cols) > 2 else ""
        if desc in ("Описание действий", "") and not loc and not ref:
            continue
        if not desc and not loc and not ref:
            continue
        rows.append({"description": desc, "location": loc, "reference": ref})
    return rows


def _merge_works_metadata(
    works: list[dict],
    descriptions: dict[int, str],
    locations: dict[int, str],
    references: dict[int, str],
    table_rows: list[dict],
) -> list[dict]:
    for work in works:
        n = work["num"]
        if not work.get("description"):
            work["description"] = descriptions.get(n, "")
        if not work.get("location"):
            work["location"] = locations.get(n, "")
        if not work.get("reference"):
            work["reference"] = references.get(n, "")

    if table_rows:
        for i, work in enumerate(works):
            if i >= len(table_rows):
                break
            row = table_rows[i]
            if not work.get("description") and row.get("description"):
                work["description"] = row["description"]
            if not work.get("location") and row.get("location"):
                work["location"] = row["location"]
            if not work.get("reference") and row.get("reference"):
                work["reference"] = row["reference"]

    for work in works:
        daily = work.get("volumes", {}).get("Объем за сутки", "")
        work["zero_daily"] = is_zero_daily_volume(daily)

    return works


def _parse_works_from_sources(
    desc_body: str,
    full_text: str,
    header_cells: dict,
    table_rows: list[dict],
) -> tuple[list[dict], str]:
    """Пробует ячейку-заголовок, затем полный текст отчёта."""
    sources: list[tuple[str, str]] = []
    if desc_body.strip():
        sources.append(("header_cell", desc_body))
    body = _extract_report_body_from_full_text(full_text)
    if body:
        sources.append(("full_text", body))

    for source, text in sources:
        part1_text, part2_text = split_part1_part2_text(text)
        part1_lines = [l for l in part1_text.splitlines() if l.strip()]
        part2_lines = [l for l in part2_text.splitlines() if l.strip()]
        works = parse_works_from_part1_lines(part1_lines)
        if works:
            descriptions = parse_numbered_items(part2_lines)
            locations: dict[int, str] = {}
            references: dict[int, str] = {}
            if "location" in header_cells:
                loc_lines = [
                    l for l in _cell_body_text(header_cells["location"]).splitlines() if l.strip()
                ]
                locations = parse_numbered_items(loc_lines)
            if "reference" in header_cells:
                ref_lines = [
                    l for l in _cell_body_text(header_cells["reference"]).splitlines() if l.strip()
                ]
                references = parse_numbered_items(ref_lines)
            works = _merge_works_metadata(works, descriptions, locations, references, table_rows)
            print(f"[SK_EXTRACT] parsed {len(works)} works from {source}")
            return works, source

    return [], "none"


def extract_sk_original(filepath: str) -> dict:
    """
    Работы из отчёта для verify.
    works[]: num, title, volumes, description, location, reference, zero_daily
    """
    try:
        doc = Document(filepath)
        header_cells, coords = find_sk_section_header_cells(doc)
        if not header_cells:
            return {"ok": False, "error": "Секция СК не найдена", "works": [], "active_works": []}

        full_text = extract_full_text(filepath)
        desc_body = _cell_body_text(header_cells["description"])
        table_rows = _extract_sk_table_rows(doc, coords) if coords else []

        works, source = _parse_works_from_sources(
            desc_body, full_text, header_cells, table_rows
        )

        if not works and table_rows:
            works = []
            for i, row in enumerate(table_rows, start=1):
                block = row.get("description", "")
                if not block:
                    continue
                part1_lines = [l for l in block.splitlines() if l.strip()]
                parsed = parse_works_from_part1_lines(part1_lines)
                if parsed:
                    works.extend(parsed)
                elif re.match(r"^\d+\.", block.strip()):
                    works.append(
                        {
                            "num": i,
                            "title": re.sub(r"^\d+\.\s*", "", block.strip()),
                            "volumes": {},
                            "description": block,
                            "location": row.get("location", ""),
                            "reference": row.get("reference", ""),
                            "zero_daily": False,
                        }
                    )
            if works:
                source = "table_rows"
                print(f"[SK_EXTRACT] parsed {len(works)} works from table_rows")

        for work in works:
            if "zero_daily" not in work:
                daily = work.get("volumes", {}).get("Объем за сутки", "")
                work["zero_daily"] = is_zero_daily_volume(daily)

        active = [w for w in works if not w.get("zero_daily")]
        print(
            f"[SK_EXTRACT] {filepath}: total={len(works)} active={len(active)} source={source}"
        )

        return {
            "ok": True,
            "works": works,
            "active_works": active,
            "source": source,
            "description_body": desc_body,
            "sk_section": extract_sk_section(filepath),
        }
    except Exception as e:
        print(f"[SK_EXTRACT] extract_sk_original error: {e}")
        return {"ok": False, "error": str(e), "works": [], "active_works": []}
