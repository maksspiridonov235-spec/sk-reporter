"""Наименование объекта с титульного листа (*_л.0_*.doc/docx)."""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
from html import unescape
from pathlib import Path
from typing import Any

from docx.oxml.ns import qn

_TITLE_PAGE_GLOBS = ("*_л.0_*.doc", "*_л.0_*.docx", "*_l.0_*.doc", "*_l.0_*.docx")
_SKIP_NAME_RE = re.compile(
    r"^(?:"
    r"экз\.?\s*№.*|"
    r"рабочая\s+документация\.?|"
    r"этап\s+строительства.*|"
    r"утверждаю.*|"
    r"согласовано.*|"
    r"разработал.*|"
    r"проверил.*"
    r")$",
    re.I,
)
_CIPHER_RE = re.compile(r"^[A-Z]{2,4}-WLL-", re.I)


def find_title_page(proj: Path, meta: dict[str, Any] | None = None) -> Path | None:
    meta = meta or {}
    explicit = (meta.get("title_page") or "").strip()
    if explicit:
        path = proj / explicit
        if path.is_file():
            return path
    for pattern in _TITLE_PAGE_GLOBS:
        matches = sorted(proj.glob(pattern))
        if matches:
            return matches[0]
    return None


def _normalize_line(text: str) -> str:
    text = unescape(text).replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _bold_texts_from_html(html: str) -> list[str]:
    out: list[str] = []
    for match in re.finditer(r"<b[^>]*>(.*?)</b>", html, re.I | re.S):
        inner = re.sub(r"<[^>]+>", "", match.group(1))
        line = _normalize_line(inner)
        if line:
            out.append(line)
    return out


def _run_is_bold(run) -> bool:
    if run.bold is True:
        return True
    r_pr = run._element.find(qn("w:rPr"))
    if r_pr is None:
        return False
    bold = r_pr.find(qn("w:b"))
    if bold is None:
        return False
    val = bold.get(qn("w:val"))
    return val is None or val not in ("0", "false")


def _bold_texts_from_docx(path: Path) -> list[str]:
    from docx import Document

    doc = Document(str(path))
    out: list[str] = []

    def collect_paragraph(para) -> None:
        parts = [run.text for run in para.runs if _run_is_bold(run) and run.text]
        if parts:
            line = _normalize_line("".join(parts))
            if line:
                out.append(line)
        elif para.runs and all(_run_is_bold(r) for r in para.runs if r.text):
            line = _normalize_line(para.text)
            if line:
                out.append(line)

    for para in doc.paragraphs:
        collect_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    collect_paragraph(para)
    return out


def _bold_texts_from_doc_mac(path: Path) -> list[str] | None:
    if sys.platform != "darwin":
        return None
    textutil = shutil.which("textutil")
    if not textutil:
        return None
    proc = subprocess.run(
        [textutil, "-convert", "html", "-stdout", str(path)],
        capture_output=True,
        text=True,
        errors="replace",
    )
    if proc.returncode != 0 or not proc.stdout.strip():
        return None
    return _bold_texts_from_html(proc.stdout)


def _pick_object_name(candidates: list[str]) -> str | None:
    parts: list[str] = []
    started = False
    for text in candidates:
        t = _normalize_line(text)
        if not t:
            continue
        if _CIPHER_RE.match(t) or re.fullmatch(r"\d{4}", t):
            break
        if _SKIP_NAME_RE.match(t) or t.upper().startswith("РАБОЧАЯ ДОКУМЕНТАЦИЯ"):
            break
        if started and re.match(r"^(?:Сети |Архитектурно|Этап строительства)", t, re.I):
            break
        if not started:
            if len(t) < 12:
                continue
            started = True
            parts.append(t)
            if re.search(r"Куст\s+скважин\s+№?\d+", t, re.I):
                break
            continue
        if re.match(r"^Куст\s+скважин", t, re.I):
            parts.append(t)
            break
        if len(t) < 40:
            parts.append(t)
            break
        break
    if parts:
        return " ".join(parts)
    for text in candidates:
        t = _normalize_line(text)
        if len(t) >= 12 and not _SKIP_NAME_RE.match(t) and not _CIPHER_RE.match(t):
            return t
    return None


def object_name_from_title_page(path: Path) -> str | None:
    suffix = path.suffix.lower()
    candidates: list[str] | None = None
    if suffix == ".docx":
        candidates = _bold_texts_from_docx(path)
    elif suffix == ".doc":
        candidates = _bold_texts_from_doc_mac(path)
    if not candidates:
        return None
    return _pick_object_name(candidates)


def resolve_object_name(proj: Path, meta: dict[str, Any] | None = None) -> tuple[str | None, str | None]:
    """(object_name, title_page_filename)"""
    meta = meta or {}
    explicit = (meta.get("object_name") or meta.get("object") or "").strip()
    if explicit:
        title_path = find_title_page(proj, meta)
        return explicit, title_path.name if title_path else None

    title_path = find_title_page(proj, meta)
    if not title_path:
        return None, None
    name = object_name_from_title_page(title_path)
    return name, title_path.name
