"""
Переключение руководителя в ежедневных отчётах (блок 2 UI).

Новые шаблоны: выпадающие списки Word (w:sdt + w:dropDownList) — должность
и ФИО в шапке и подвале. Старые шаблоны: rule-based замена в ячейках таблицы.
"""

from __future__ import annotations

import re
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell
from lxml import etree

from sk_reporter.docx_processing import FONT_NAME, FONT_SIZE_HALF_POINTS
from sk_reporter.template_layout import _main_table_indices

_FONT_RPR_TAGS = ("rFonts", "sz", "szCs")
_THEME_ATTR_MARKERS = ("theme", "Theme")

LeaderId = Literal["aniskov", "mandzhiev"]

LEADER_TARGETS: dict[LeaderId, dict[str, str]] = {
    "aniskov": {
        "header_title": "Руководитель",
        "header_fio": "Аниськов Владимир Иванович",
        "footer_role": "Руководитель проекта СК",
        "footer_fio": "Аниськов Владимир Иванович",
    },
    "mandzhiev": {
        "header_title": "И.О. Руководителя",
        "header_fio": "Манджиев Игорь Александрович",
        "footer_role": "И.О. Руководителя проекта СК",
        "footer_fio": "Манджиев Игорь Александрович",
    },
}

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = f"{{{_W_NS}}}"

_HEADER_TITLE_OPTIONS = frozenset({"Руководитель", "И.О. Руководителя"})
_FOOTER_ROLE_OPTIONS = frozenset(
    {"Руководитель проекта СК", "И.О. Руководителя проекта СК"}
)
_LEADER_FIO_OPTIONS = frozenset(
    {"Аниськов Владимир Иванович", "Манджиев Игорь Александрович"}
)

# Варианты должности в подвале (для замены в run'ах)
_FOOTER_ROLE_VARIANTS = [
    "И.О. Руководителя проекта СК",
    "И.о. Руководителя проекта СК",
    "И.о. руководителя проекта СК",
    "Руководитель проекта СК",
    "Руководитель проекта",
]

_HEADER_TITLE_VARIANTS = [
    "И.О. Руководителя",
    "И.о. Руководителя",
    "И.о. руководителя",
    "Руководитель проекта СК",
    "Руководитель проекта",
    "Руководитель",
]

_FIO_REGEXES = [
    re.compile(r"Аниськов\s+В(?:ладимир)?\s+И(?:ванович)?\.?", re.I),
    re.compile(r"Манджиев\s+И(?:горь)?\s+А(?:лександрович)?\.?", re.I),
    re.compile(r"Аниськов\s+В\.?\s*И\.?", re.I),
    re.compile(r"Манджиев\s+И\.?\s*А\.?", re.I),
]

# Должность в подвале (в т.ч. разбитая по run/абзацам)
_FOOTER_ROLE_RE = re.compile(
    r"(?:и\.?\s*о\.?\s*)?руководител\w*\s+проект\w*(?:\s+с\.?\s*к\.?)?",
    re.I,
)

HEADER_ZONE_ROWS = 14
FOOTER_ZONE_ROWS = 25


@dataclass(frozen=True)
class CellSlot:
    row: int
    col: int


@dataclass
class LeaderSlots:
    header_title: CellSlot
    header_fio: CellSlot
    footer_role: CellSlot
    footer_fio: CellSlot


def _norm(text: str) -> str:
    cleaned = text.replace("\xa0", " ").replace("\u200b", "").replace("\ufeff", "")
    return " ".join(cleaned.split()).strip()


def _sdt_dropdown_options(sdt) -> frozenset[str]:
    opts: set[str] = set()
    for li in sdt.findall(f".//{_W}listItem"):
        val = li.get(f"{_W}displayText") or li.get(f"{_W}value")
        if val:
            opts.add(val.strip())
    return frozenset(opts)


def _sdt_text(sdt) -> str:
    ts = sdt.findall(f".//{_W}sdtContent//{_W}t")
    return _norm("".join(t.text or "" for t in ts))


def _set_sdt_text(sdt, new_text: str) -> bool:
    if _sdt_text(sdt) == _norm(new_text):
        return False
    ts = sdt.findall(f".//{_W}sdtContent//{_W}t")
    if not ts:
        return False
    updated = False
    for t in ts:
        if t.text and not updated:
            t.text = new_text
            updated = True
        elif t.text:
            t.text = ""
    if not updated:
        ts[0].text = new_text
    return True


def _classify_leader_sdt(sdt) -> str | None:
    """header_title | footer_role | fio — по набору пунктов выпадающего списка."""
    opts = _sdt_dropdown_options(sdt)
    if not opts:
        return None
    if opts == _LEADER_FIO_OPTIONS:
        return "fio"
    if opts == _FOOTER_ROLE_OPTIONS:
        return "footer_role"
    if opts == _HEADER_TITLE_OPTIONS:
        return "header_title"
    joined = " ".join(opts).lower()
    if "аниськов" in joined and "манджиев" in joined:
        return "fio"
    if any("проекта" in o.lower() for o in opts) and any(
        "руководител" in o.lower() for o in opts
    ):
        return "footer_role"
    if any("руководител" in o.lower() for o in opts) and not any(
        "проект" in o.lower() for o in opts
    ):
        return "header_title"
    return None


def _apply_leader_dropdowns(doc: Document, leader: LeaderId) -> int | None:
    """
    Переключение через выпадающие списки Word.
    None — в документе нет leader-SDT; иначе число изменённых полей.
    """
    targets = LEADER_TARGETS[leader]
    sdts = doc.element.body.findall(f".//{_W}sdt")
    header_title_sdts: list = []
    footer_role_sdts: list = []
    fio_sdts: list = []

    for sdt in sdts:
        kind = _classify_leader_sdt(sdt)
        if kind == "header_title":
            header_title_sdts.append(sdt)
        elif kind == "footer_role":
            footer_role_sdts.append(sdt)
        elif kind == "fio":
            fio_sdts.append(sdt)

    if not header_title_sdts and not footer_role_sdts and not fio_sdts:
        return None

    changes = 0
    for sdt in header_title_sdts:
        if _set_sdt_text(sdt, targets["header_title"]):
            changes += 1
    for sdt in footer_role_sdts:
        if _set_sdt_text(sdt, targets["footer_role"]):
            changes += 1
    fio_target = targets["header_fio"]
    for sdt in fio_sdts:
        if _set_sdt_text(sdt, fio_target):
            changes += 1
    return changes


def _cell_text(cell: _Cell) -> str:
    return _norm(cell.text)


def _rPr_ensure_bold_flags(rPr) -> bool:
    """w:b и w:bCs на уровне rPr (как в шаблоне Word для кириллицы)."""
    changed = False
    for local in ("b", "bCs"):
        tag = qn(f"w:{local}")
        el = rPr.find(tag)
        if el is not None:
            val = el.get(qn("w:val"))
            if val in ("0", "false", "off"):
                el.attrib.pop(qn("w:val"), None)
                changed = True
        else:
            etree.SubElement(rPr, tag)
            changed = True
    return changed


def _rPr_strip_theme_fonts(rFonts) -> bool:
    """Убрать asciiTheme/minorHAnsi — иначе Word показывает Calibri поверх ascii."""
    changed = False
    for attr in list(rFonts.attrib):
        local = attr.split("}")[-1] if "}" in attr else attr
        if any(m in local for m in _THEME_ATTR_MARKERS):
            del rFonts.attrib[attr]
            changed = True
    return changed


def _rPr_ensure_times_font(rPr) -> bool:
    """Times New Roman 10 pt, если в rPr нет шрифта (как в prepare/format_fonts)."""
    changed = False
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
        changed = True
    if _rPr_strip_theme_fonts(rFonts):
        changed = True
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        key = qn(f"w:{attr}")
        if rFonts.get(key) != FONT_NAME:
            rFonts.set(key, FONT_NAME)
            changed = True
    for local in ("sz", "szCs"):
        tag = qn(f"w:{local}")
        el = rPr.find(tag)
        if el is None:
            el = etree.SubElement(rPr, tag)
            changed = True
        if el.get(qn("w:val")) != FONT_SIZE_HALF_POINTS:
            el.set(qn("w:val"), FONT_SIZE_HALF_POINTS)
            changed = True
    return changed


def _rPr_copy_font_from_template(template, rPr) -> bool:
    """Копирует rFonts/sz/szCs из pPr/rPr в run, чтобы не сбить Times New Roman."""
    if template is None:
        return False
    changed = False
    for tag_local in _FONT_RPR_TAGS:
        tag = qn(f"w:{tag_local}")
        src = template.find(tag)
        if src is None:
            continue
        dest = rPr.find(tag)
        if dest is None:
            rPr.append(deepcopy(src))
            changed = True
        elif tag_local == "rFonts":
            for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                key = qn(f"w:{attr}")
                val = src.get(key)
                if val and dest.get(key) != val:
                    dest.set(key, val)
                    changed = True
        else:
            val = src.get(qn("w:val"))
            if val and dest.get(qn("w:val")) != val:
                dest.set(qn("w:val"), val)
                changed = True
    return changed


def _paragraph_default_rPr(para):
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        return None
    return pPr.find(qn("w:rPr"))


def _capture_format_rPr(para):
    """Образец форматирования: pPr/rPr или первый run с rFonts."""
    template = _paragraph_default_rPr(para)
    if template is not None and template.find(qn("w:rFonts")) is not None:
        return template
    for run in para.runs:
        rPr = run._r.find(qn("w:rPr"))
        if rPr is not None and rPr.find(qn("w:rFonts")) is not None:
            return rPr
    return template


def _rewrite_paragraph_single_run(para, text: str, *, bold: bool) -> bool:
    """
    Один run с явным Times New Roman (без para.text — он ломает шрифт/жирный).
    """
    new_text = text
    if _norm(para.text) == _norm(new_text):
        changed = False
        if bold:
            changed = _ensure_paragraph_role_bold(para)
        return changed

    template = _capture_format_rPr(para)
    p = para._p
    for child in list(p):
        if child.tag == qn("w:r"):
            p.remove(child)

    r = etree.SubElement(p, qn("w:r"))
    rPr = etree.SubElement(r, qn("w:rPr"))
    if template is not None:
        for child in template:
            if child.tag in (qn("w:b"), qn("w:bCs")):
                continue
            rPr.append(deepcopy(child))
    _rPr_ensure_times_font(rPr)
    if bold:
        _rPr_ensure_bold_flags(rPr)

    t = etree.SubElement(r, qn("w:t"))
    if new_text.startswith(" ") or new_text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = new_text

    pPr = p.find(qn("w:pPr"))
    if pPr is None:
        pPr = para._p.get_or_add_pPr()
    rPr_p = pPr.find(qn("w:rPr"))
    if rPr_p is None:
        rPr_p = etree.SubElement(pPr, qn("w:rPr"))
    _rPr_ensure_times_font(rPr_p)
    if bold:
        _rPr_ensure_bold_flags(rPr_p)
    return True


def _text_is_target_role(text: str, leader: LeaderId) -> bool:
    t = _norm(text)
    if not t:
        return False
    targets = LEADER_TARGETS[leader]
    if t == _norm(targets["header_title"]) or t == _norm(targets["footer_role"]):
        return True
    if _is_header_title_row_cell(t) or _is_footer_role(t) or _FOOTER_ROLE_RE.search(t):
        return True
    return any(_norm(v) == t for v in _HEADER_TITLE_VARIANTS + _FOOTER_ROLE_VARIANTS)


def _ensure_paragraph_role_bold(para) -> bool:
    """Жирный + Times New Roman: pPr/rPr и каждый run (без run.bold — ломает шрифт)."""
    changed = False
    pPr = para._p.find(qn("w:pPr"))
    if pPr is None:
        pPr = para._p.get_or_add_pPr()
    rPr_p = pPr.find(qn("w:rPr"))
    if rPr_p is None:
        rPr_p = etree.SubElement(pPr, qn("w:rPr"))
    if _rPr_ensure_times_font(rPr_p):
        changed = True
    if _rPr_ensure_bold_flags(rPr_p):
        changed = True

    for run in para.runs:
        r = run._r
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = etree.SubElement(r, qn("w:rPr"))
        if _rPr_copy_font_from_template(rPr_p, rPr):
            changed = True
        if _rPr_ensure_times_font(rPr):
            changed = True
        if _rPr_ensure_bold_flags(rPr):
            changed = True
    return changed


def _ensure_runs_bold(para) -> bool:
    return _ensure_paragraph_role_bold(para)


def _set_paragraph_text(para, text: str, *, bold: bool = False) -> bool:
    if bold:
        return _rewrite_paragraph_single_run(para, text, bold=True)
    para.text = text
    return True


def _finalize_role_cell_format(cell: _Cell, leader: LeaderId, *, is_header: bool) -> bool:
    targets = LEADER_TARGETS[leader]
    target = targets["header_title"] if is_header else targets["footer_role"]
    changed = False
    for para in cell.paragraphs:
        ptext = _norm(para.text)
        if not ptext or "лицеванов" in ptext.lower():
            continue
        if is_header:
            if not (_is_header_title_row_cell(ptext) or ptext == _norm(target)):
                continue
        elif not (_cell_has_footer_role_text(ptext) or ptext == _norm(target)):
            continue
        if _rewrite_paragraph_single_run(para, target, bold=True):
            changed = True
    return changed


def _apply_signature_role_bold(
    table,
    leader: LeaderId,
    header_row: int = 0,
    footer_role_row: int | None = None,
) -> int:
    """Финальный проход: TNR + жирный во всех ячейках должности (шапка + подвал)."""
    nrows = len(table.rows)
    changes = 0
    seen_tc: set[int] = set()

    for ri in range(min(HEADER_ZONE_ROWS, nrows)):
        for _ci, cell in _zone_cells(table.rows[ri]):
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            text = _cell_text(cell)
            if not (_is_header_title_row_cell(text) or _text_is_target_role(text, leader)):
                continue
            seen_tc.add(tc_id)
            if _finalize_role_cell_format(cell, leader, is_header=True):
                changes += 1

    seen_tc.clear()
    scan_from = _footer_scan_from(nrows, header_row, footer_role_row)
    for ri in range(scan_from, nrows):
        for _ci, cell in _zone_cells(table.rows[ri]):
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            text = _cell_text(cell)
            if not (_cell_has_footer_role_text(text) or _text_is_target_role(text, leader)):
                continue
            seen_tc.add(tc_id)
            if _finalize_role_cell_format(cell, leader, is_header=False):
                changes += 1

    return changes


def _is_leader_fio(text: str) -> bool:
    t = text.lower()
    if "аниськов" in t or "манджиев" in t:
        return True
    return any(r.search(text) for r in _FIO_REGEXES)


def _is_header_title(text: str) -> bool:
    """Короткая должность в шапке (без «проекта»)."""
    t = text.lower()
    if "руководител" not in t:
        return False
    if "лицеванов" in t or "проекта" in t:
        return False
    return True


def _is_header_title_row_cell(text: str) -> bool:
    """Ячейка должности в строке шапки (в т.ч. ошибочно «…проекта СК» в GRID_COLS_7)."""
    t = text.lower()
    if "руководител" not in t:
        return False
    if "лицеванов" in t:
        return False
    if _is_leader_fio(text) and not _FOOTER_ROLE_RE.search(text):
        return False
    return True


def _is_footer_role(text: str) -> bool:
    t = text.lower()
    if "лицеванов" in t:
        return False
    if _FOOTER_ROLE_RE.search(t):
        return True
    if "руководител" in t and "проект" in t:
        return True
    # «Руководитель проекта СК» иногда без слова «проект» в отдельной ячейке/run
    return "руководител" in t and re.search(r"(?:^|\s)с\.?\s*к\.?(?:\s|$)", t, re.I) is not None


def _is_footer_role_paragraph(text: str) -> bool:
    """Абзац с полной должностью (не чистое ФИО)."""
    t = _norm(text)
    if not t or "лицеванов" in t.lower():
        return False
    if _is_leader_fio(t) and not _FOOTER_ROLE_RE.search(t):
        return False
    return _is_footer_role(t)


def _is_fio_paragraph(text: str) -> bool:
    t = _norm(text)
    return bool(t) and _is_leader_fio(t) and not _is_footer_role_paragraph(t)


def _skip_cell(text: str) -> bool:
    t = text.lower()
    return not text or "лицеванов" in t


def _zone_cells(row) -> list[tuple[int, _Cell]]:
    """Все непустые уникальные ячейки строки (для подвала — без обрезки по col)."""
    out: list[tuple[int, _Cell]] = []
    seen: set[int] = set()
    for ci, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        if _cell_text(cell):
            out.append((ci, cell))
    return out


def _right_band_cells(row, min_col: int = 2) -> list[tuple[int, _Cell]]:
    """Уникальные непустые ячейки правой части строки (подписи)."""
    cells = list(enumerate(row.cells))
    if not cells:
        return []
    start = max(min_col, len(cells) - 4)
    out: list[tuple[int, _Cell]] = []
    seen: set[int] = set()
    for ci, cell in cells:
        if ci < start:
            continue
        tc_id = id(cell._tc)
        if tc_id in seen:
            continue
        seen.add(tc_id)
        if _cell_text(cell):
            out.append((ci, cell))
    return out


def find_leader_slots(doc: Document) -> LeaderSlots | None:
    """Диагностика: первая найденная четвёрка слотов (как раньше)."""
    indices = _main_table_indices(doc)
    if not indices:
        return None
    table = doc.tables[indices[0]]
    nrows = len(table.rows)
    if nrows < 12:
        return None

    header_title: CellSlot | None = None
    header_fio: CellSlot | None = None
    for ri in range(min(HEADER_ZONE_ROWS, nrows)):
        titles: list[int] = []
        fios: list[int] = []
        for ci, cell in _right_band_cells(table.rows[ri]):
            text = _cell_text(cell)
            if _is_leader_fio(text):
                fios.append(ci)
            elif _is_header_title_row_cell(text):
                titles.append(ci)
        if titles and fios:
            header_title = CellSlot(ri, titles[-1])
            header_fio = CellSlot(ri, fios[-1])
            break

    if header_title is None or header_fio is None:
        return None

    footer_role: CellSlot | None = None
    footer_fio: CellSlot | None = None
    scan_from = max(nrows - FOOTER_ZONE_ROWS, header_title.row + 1)
    for ri in range(nrows - 1, scan_from - 1, -1):
        for ci, cell in _right_band_cells(table.rows[ri]):
            text = _cell_text(cell)
            if footer_fio is None and _is_leader_fio(text):
                footer_fio = CellSlot(ri, ci)
            if footer_role is None and _is_footer_role(text):
                footer_role = CellSlot(ri, ci)

    if footer_role is None or footer_fio is None:
        return None

    return LeaderSlots(
        header_title=header_title,
        header_fio=header_fio,
        footer_role=footer_role,
        footer_fio=footer_fio,
    )


def _write_cell_text(cell: _Cell, new_text: str, *, bold: bool = False) -> bool:
    if _cell_text(cell) == new_text:
        if bold and cell.paragraphs:
            fmt_changed = False
            for para in cell.paragraphs:
                if _norm(para.text) and _ensure_runs_bold(para):
                    fmt_changed = True
            return fmt_changed
        return False
    if cell.paragraphs:
        _set_paragraph_text(cell.paragraphs[0], new_text, bold=bold)
        for para in cell.paragraphs[1:]:
            para.text = ""
    else:
        cell.text = new_text
    return True


def _replace_in_runs_paragraph(
    para, old_variants: list[str], new_text: str, *, bold: bool = False
) -> bool:
    changed = False
    for run in para.runs:
        original = run.text
        updated = original
        for old in old_variants:
            if old in updated:
                updated = updated.replace(old, new_text)
            else:
                m = re.search(re.escape(old), updated, re.I)
                if m:
                    updated = updated[: m.start()] + new_text + updated[m.end() :]
        if updated != original:
            run.text = updated
            changed = True
    if changed and bold:
        _ensure_runs_bold(para)
    return changed


def _replace_in_runs(
    cell: _Cell, old_variants: list[str], new_text: str, *, bold: bool = False
) -> bool:
    changed = False
    for para in cell.paragraphs:
        if _replace_in_runs_paragraph(para, old_variants, new_text, bold=bold):
            changed = True
    return changed


def _replace_footer_role_in_paragraph(para, target: str) -> bool:
    ptext = _norm(para.text)
    if not ptext or "лицеванов" in ptext.lower():
        return False
    if not _is_footer_role_paragraph(ptext):
        return False
    if ptext == target:
        return _ensure_runs_bold(para)
    if _replace_in_runs_paragraph(para, _FOOTER_ROLE_VARIANTS, target, bold=True):
        return True
    new_pt = _FOOTER_ROLE_RE.sub(target, para.text, count=1)
    if new_pt != para.text:
        _set_paragraph_text(para, new_pt, bold=True)
        return True
    if ptext != target:
        _set_paragraph_text(para, target, bold=True)
        return True
    return _ensure_runs_bold(para)


def _replace_title_in_cell(cell: _Cell, target: str) -> bool:
    if not _is_header_title_row_cell(_cell_text(cell)) and _cell_text(cell) != target:
        return False
    if not cell.paragraphs:
        return False
    return _rewrite_paragraph_single_run(cell.paragraphs[0], target, bold=True)


def _replace_footer_role_in_cell(cell: _Cell, target: str) -> bool:
    if _cell_text(cell) == target:
        changed = False
        for para in cell.paragraphs:
            if _is_footer_role_paragraph(para.text) and _ensure_runs_bold(para):
                changed = True
        return changed
    changed = False
    for para in cell.paragraphs:
        if _replace_footer_role_in_paragraph(para, target):
            changed = True
    if changed:
        return True
    if _is_footer_role(_cell_text(cell)):
        return _write_cell_text(cell, target, bold=True)
    return False


def _replace_fio_in_paragraph(para, target_fio: str) -> bool:
    if _norm(para.text) == target_fio:
        return False
    changed = False
    for run in para.runs:
        new = run.text
        for rx in _FIO_REGEXES:
            new = rx.sub(target_fio, new)
        if new != run.text:
            run.text = new
            changed = True
    if changed:
        return True
    if _is_leader_fio(_norm(para.text)):
        para.text = target_fio
        return True
    return False


def _replace_fio_in_cell(cell: _Cell, target_fio: str) -> bool:
    if _cell_text(cell) == target_fio:
        return False
    changed = False
    for para in cell.paragraphs:
        if _replace_fio_in_paragraph(para, target_fio):
            changed = True
    if changed:
        return True
    if _is_leader_fio(_cell_text(cell)) or not _cell_text(cell):
        return _write_cell_text(cell, target_fio)
    return False


def _apply_header_zone(table, leader: LeaderId) -> int:
    targets = LEADER_TARGETS[leader]
    nrows = len(table.rows)
    changes = 0
    seen_tc: set[int] = set()

    for ri in range(min(HEADER_ZONE_ROWS, nrows)):
        for _ci, cell in _right_band_cells(table.rows[ri]):
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            text = _cell_text(cell)
            if _skip_cell(text):
                continue

            if _is_leader_fio(text):
                seen_tc.add(tc_id)
                if _replace_fio_in_cell(cell, targets["header_fio"]):
                    changes += 1
            elif _is_header_title_row_cell(text):
                seen_tc.add(tc_id)
                if _replace_title_in_cell(cell, targets["header_title"]):
                    changes += 1

    return changes


def _footer_scan_from(nrows: int, header_row: int, footer_role_row: int | None = None) -> int:
    """Нижняя часть таблицы: не раньше шапки и не уже 25 последних строк."""
    scan_from = max(header_row + 10, nrows - FOOTER_ZONE_ROWS, HEADER_ZONE_ROWS)
    if footer_role_row is not None:
        scan_from = min(scan_from, max(HEADER_ZONE_ROWS, footer_role_row - 3))
    return scan_from


def _cell_looks_like_footer_role_fragment(text: str) -> bool:
    if _skip_cell(text):
        return False
    if _is_leader_fio(text) and not _FOOTER_ROLE_RE.search(text):
        return False
    tl = text.lower()
    if "руководител" in tl or "проект" in tl or _FOOTER_ROLE_RE.search(text):
        return True
    return bool(re.search(r"и\.?\s*о\.?", tl, re.I) and ("проект" in tl or "руководител" in tl))


def _row_footer_role_combined_text(row) -> str:
    parts: list[str] = []
    for _ci, cell in _zone_cells(row):
        text = _cell_text(cell)
        if _cell_looks_like_footer_role_fragment(text):
            parts.append(text)
    return _norm(" ".join(parts))


def _apply_footer_role_row_combined(row, target: str) -> bool:
    """Должность в подвале, разбитая по нескольким ячейкам одной строки."""
    combined = _row_footer_role_combined_text(row)
    if not combined or not (_is_footer_role(combined) or _FOOTER_ROLE_RE.search(combined)):
        return False
    if _norm(combined) == _norm(target):
        return False

    frag: list[tuple[_Cell, str]] = []
    for _ci, cell in _zone_cells(row):
        text = _cell_text(cell)
        if _cell_looks_like_footer_role_fragment(text):
            frag.append((cell, text))
    if not frag:
        return False

    primary = max(frag, key=lambda x: len(x[1]))[0]
    primary_tc = id(primary._tc)
    changed = _write_cell_text(primary, target, bold=True)
    for cell, text in frag:
        if cell is primary or id(cell._tc) == primary_tc:
            continue
        if text and _write_cell_text(cell, ""):
            changed = True
    return changed


def _cell_has_footer_role_text(text: str) -> bool:
    if _skip_cell(text):
        return False
    if _is_leader_fio(text) and not _FOOTER_ROLE_RE.search(text):
        return False
    return bool(_is_footer_role(text) or _FOOTER_ROLE_RE.search(text))


def _force_replace_footer_role_in_cell(cell: _Cell, target: str) -> bool:
    """Замена должности в подвале одним run (TNR + жирный)."""
    changed = False
    for para in cell.paragraphs:
        raw = para.text
        pnorm = _norm(raw)
        if not pnorm or "лицеванов" in pnorm.lower():
            continue
        if _is_fio_paragraph(raw):
            continue
        if not (_is_footer_role_paragraph(raw) or _FOOTER_ROLE_RE.search(raw)):
            continue
        if _rewrite_paragraph_single_run(para, target, bold=True):
            changed = True

    if changed:
        return True

    full = _cell_text(cell)
    if not _cell_has_footer_role_text(full):
        return False
    if cell.paragraphs and _rewrite_paragraph_single_run(
        cell.paragraphs[0], target, bold=True
    ):
        for para in cell.paragraphs[1:]:
            para.text = ""
        return True
    return False


def _apply_footer_zone(
    table,
    leader: LeaderId,
    header_row: int = 0,
    footer_role_row: int | None = None,
) -> int:
    targets = LEADER_TARGETS[leader]
    nrows = len(table.rows)
    scan_from = _footer_scan_from(nrows, header_row, footer_role_row)
    changes = 0

    for ri in range(scan_from, nrows):
        row = table.rows[ri]
        row_changed = False

        for _ci, cell in _zone_cells(row):
            for para in cell.paragraphs:
                ptext = _norm(para.text)
                if not ptext or "лицеванов" in ptext.lower():
                    continue
                if _is_fio_paragraph(ptext):
                    if _replace_fio_in_paragraph(para, targets["footer_fio"]):
                        row_changed = True

            text = _cell_text(cell)
            if _cell_has_footer_role_text(text):
                if _force_replace_footer_role_in_cell(cell, targets["footer_role"]):
                    row_changed = True
            elif _is_leader_fio(text):
                if _replace_fio_in_cell(cell, targets["footer_fio"]):
                    row_changed = True

        if _apply_footer_role_row_combined(row, targets["footer_role"]):
            row_changed = True

        if row_changed:
            changes += 1

    return changes


def switch_leader_in_docx(filepath: str, leader: LeaderId) -> tuple[bool, str, int]:
    path = Path(filepath)
    try:
        from sk_reporter.docx_processing import remove_editing_restrictions

        remove_editing_restrictions(str(path))
        doc = Document(str(path))

        dd_changes = _apply_leader_dropdowns(doc, leader)
        if dd_changes is not None:
            if dd_changes == 0:
                return True, "уже нужный руководитель", 0
            doc.save(str(path))
            return True, f"замен: {dd_changes}", dd_changes

        indices = _main_table_indices(doc)
        if not indices:
            return False, "нет таблиц", 0
        table = doc.tables[indices[0]]

        slots = find_leader_slots(doc)
        header_row = slots.header_title.row if slots else 0
        footer_role_row = slots.footer_role.row if slots else None
        changes = _apply_header_zone(table, leader) + _apply_footer_zone(
            table, leader, header_row, footer_role_row
        )
        changes += _apply_signature_role_bold(
            table, leader, header_row, footer_role_row
        )

        if changes == 0:
            if slots is None:
                return False, "не найдены ячейки руководителя", 0
            return True, "уже нужный руководитель", 0

        doc.save(str(path))
        return True, f"замен: {changes}", changes
    except Exception as e:
        return False, str(e), 0


def switch_leader_batch(filepaths: list[str], leader: LeaderId) -> list[dict]:
    results = []
    for fp in filepaths:
        name = Path(fp).name
        ok, msg, n = switch_leader_in_docx(fp, leader)
        results.append({"filename": name, "ok": ok, "msg": msg, "changes": n})
    return results
