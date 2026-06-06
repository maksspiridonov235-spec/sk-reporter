"""Парсинг ведомости объёмов работ (ВОР) из docx."""

from __future__ import annotations

import json
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document

_STAGE_RE = re.compile(r"^\d+\s*этап", re.I)


@dataclass
class VorWorkItem:
    name: str
    unit: str = ""
    quantity: str = ""
    note: str = ""


@dataclass
class VorObject:
    title: str
    works: list[VorWorkItem] = field(default_factory=list)


@dataclass
class VorStage:
    title: str
    objects: list[VorObject] = field(default_factory=list)
    works: list[VorWorkItem] = field(default_factory=list)


@dataclass
class VorDocument:
    project_id: str
    source: str
    stages: list[VorStage] = field(default_factory=list)


def _row_texts(row) -> list[str]:
    return [c.text.strip().replace("\n", " ") for c in row.cells]


def _is_stage_row(cells: list[str]) -> bool:
    if not cells or not cells[0]:
        return False
    return _STAGE_RE.match(cells[0]) is not None


def _is_object_row(cells: list[str]) -> bool:
    if not cells or not cells[0]:
        return False
    name = cells[0]
    if _STAGE_RE.match(name):
        return False
    unique = {c for c in cells if c}
    return len(unique) == 1


def parse_vor_docx(path: Path, project_id: str = "") -> VorDocument:
    doc = Document(str(path))
    if not doc.tables:
        raise ValueError(f"Нет таблиц в {path.name}")

    table = doc.tables[0]
    result = VorDocument(project_id=project_id or path.parent.name, source=path.name)

    current_stage: Optional[VorStage] = None
    current_object: Optional[VorObject] = None

    for row in table.rows[1:]:
        cells = _row_texts(row)
        if not any(cells):
            continue

        if _is_stage_row(cells):
            current_stage = VorStage(title=cells[0])
            result.stages.append(current_stage)
            current_object = None
            continue

        if _is_object_row(cells):
            if current_stage is None:
                current_stage = VorStage(title="(без этапа)")
                result.stages.append(current_stage)
            current_object = VorObject(title=cells[0])
            current_stage.objects.append(current_object)
            continue

        name = cells[1] if len(cells) > 1 else cells[0]
        if not name:
            continue

        item = VorWorkItem(
            name=name,
            unit=cells[2] if len(cells) > 2 else "",
            quantity=cells[3] if len(cells) > 3 else "",
            note=cells[4] if len(cells) > 4 else "",
        )
        if current_object is not None:
            current_object.works.append(item)
        elif current_stage is not None:
            current_stage.works.append(item)

    return result


def vor_to_dict(vor: VorDocument) -> dict:
    def stage_dict(stage: VorStage) -> dict:
        return {
            "title": stage.title,
            "objects": [
                {"title": o.title, "works": [asdict(w) for w in o.works]}
                for o in stage.objects
            ],
            "works": [asdict(w) for w in stage.works],
        }

    return {
        "project_id": vor.project_id,
        "source": vor.source,
        "stages": [stage_dict(s) for s in vor.stages],
    }


def write_vor_cache(vor_path: Path, cache_path: Optional[Path] = None) -> Path:
    vor = parse_vor_docx(vor_path, project_id=vor_path.parent.name)
    out = cache_path or vor_path.parent / "vor.json"
    out.write_text(json.dumps(vor_to_dict(vor), ensure_ascii=False, indent=2), encoding="utf-8")
    return out
