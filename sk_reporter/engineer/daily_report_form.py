"""HTML-форма ежедневного отчёта — сетка читается из report_template.docx."""

from __future__ import annotations

from dataclasses import dataclass
from html import escape
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

_TEMPLATE = Path(__file__).resolve().parents[2] / "data" / "engineer" / "report_template.docx"

_STATIC = {
    "ООО «СПД»",
    "123242, г. Москва, Новинский бульвар, 31, 6 этаж",
    "8(495)5189720",
    "info@spd.ru",
    "MOS/25/0008",
    "01.07.2025",
    "1",
    "Согласно списку рассылки",
}


@dataclass(frozen=True)
class TableLayout:
    """Геометрия главной таблицы — только из .docx."""

    table_width_dxa: int
    col_widths_dxa: tuple[int, ...]
    row_heights_dxa: tuple[int | None, ...]
    row_height_rules: tuple[str | None, ...]


def dxa_to_px(dxa: int) -> float:
    """DXA (twips) → CSS px при 96 dpi: 1440 dxa = 1 inch."""
    return dxa / 15.0


def read_layout_from_docx(path: Path | None = None) -> TableLayout:
    """Читает tblW, tblGrid и trHeight из шаблона пользователя."""
    doc_path = path or _TEMPLATE
    doc = Document(str(doc_path))
    table = doc.tables[0]
    tbl = table._tbl

    tbl_pr = tbl.tblPr
    table_width_dxa = 0
    if tbl_pr is not None:
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is not None and tbl_w.get(qn("w:type")) == "dxa":
            table_width_dxa = int(tbl_w.get(qn("w:w")) or 0)

    grid = tbl.tblGrid
    col_widths: list[int] = []
    if grid is not None:
        for grid_col in grid.findall(qn("w:gridCol")):
            col_widths.append(int(grid_col.get(qn("w:w")) or 0))

    if not col_widths:
        raise ValueError(f"В {doc_path} нет tblGrid — не из чего взять сетку")

    if table_width_dxa <= 0:
        table_width_dxa = sum(col_widths)

    row_heights: list[int | None] = []
    row_rules: list[str | None] = []
    for row in table.rows:
        tr_pr = row._tr.trPr
        height: int | None = None
        rule: str | None = None
        if tr_pr is not None:
            tr_height = tr_pr.find(qn("w:trHeight"))
            if tr_height is not None:
                raw = tr_height.get(qn("w:val"))
                if raw:
                    height = int(raw)
                rule = tr_height.get(qn("w:hRule"))
        row_heights.append(height)
        row_rules.append(rule)

    return TableLayout(
        table_width_dxa=table_width_dxa,
        col_widths_dxa=tuple(col_widths),
        row_heights_dxa=tuple(row_heights),
        row_height_rules=tuple(row_rules),
    )


_ROW_TALL = frozenset({16, 19, 30, 31, 33, 34, 35, 36, 37})
_ROW_SK = frozenset({22, 23, 24, 27, 28})


def render_layout_css(layout: TableLayout) -> str:
    """Пропорции колонок из .docx; таблица на всю ширину карточки."""
    total_dxa = layout.table_width_dxa or sum(layout.col_widths_dxa)
    col_rules = "".join(
        f".dr-table col.dr-c{i} {{ width: {w / total_dxa * 100:.6f}%; }}"
        for i, w in enumerate(layout.col_widths_dxa)
    )
    row_rules: list[str] = []
    for i, (h_dxa, h_rule) in enumerate(
        zip(layout.row_heights_dxa, layout.row_height_rules, strict=True)
    ):
        if not h_dxa or i in _ROW_TALL or i in _ROW_SK or i == 17:
            continue
        h_px = dxa_to_px(h_dxa)
        minmax = "min-height" if h_rule in (None, "atLeast") else "height"
        row_rules.append(f".dr-table tr.dr-r{i} {{ {minmax}: {h_px:.4f}px; }}")
    rows_css = "\n".join(row_rules)
    return f"""
.dr-table {{
  width: 100%;
  max-width: 100%;
}}
{col_rules}
{rows_css}
""".strip()


def _grid_rows(table) -> list[list[dict | str | None]]:
    rows: list[list[dict | str | None]] = []
    for row in table.rows:
        grid: list[dict | str | None] = [None] * 6
        col = 0
        seen: set[int] = set()
        for cell in row.cells:
            tc = cell._tc
            if id(tc) in seen:
                continue
            seen.add(id(tc))
            tc_pr = tc.tcPr
            if tc_pr is not None:
                vm = tc_pr.find(qn("w:vMerge"))
                if vm is not None and vm.get(qn("w:val")) == "continue":
                    continue
            while col < 6 and grid[col] == "SPAN":
                col += 1
            if col >= 6:
                break
            colspan = 1
            if tc_pr is not None:
                g = tc_pr.find(qn("w:gridSpan"))
                if g is not None:
                    colspan = int(g.get(qn("w:val")))
            grid[col] = {"colspan": colspan, "text": cell.text.strip(), "col": col}
            for j in range(1, colspan):
                if col + j < 6:
                    grid[col + j] = "SPAN"
            col += colspan
        rows.append(grid)
    return rows


def _is_section(text: str) -> bool:
    line = text.split("\n")[0].strip()
    return line.isupper() and len(line) > 10


def _cell_content(ri: int, cell: dict) -> tuple[str, str]:
    text = cell["text"]
    col = cell["col"]
    cs = cell["colspan"]

    if ri == 0 and col == 0:
        return "dr-title", text.replace("\n", "<br>")
    if text == "Страница":
        return "dr-label dr-center", "Страница"
    if ri == 0 and text == "1":
        return "dr-center", "1"
    if text in _STATIC:
        return "dr-static", escape(text)
    if _is_section(text):
        return "dr-section", escape(text).replace("\n", "<br>")

    if ri == 1 and col == 1:
        return "", '<input type="text" name="report_no" class="dr-input"/>'
    if ri == 1 and col == 3:
        return "", '<input type="text" name="object" class="dr-input"/>'
    if ri == 2 and col == 1:
        return "", '<input type="date" name="report_date" id="reportDate" class="dr-input"/>'
    if ri == 2 and col == 3:
        return "", '<input type="text" name="control_direction" class="dr-input"/>'
    if ri == 6 and col == 1:
        return "", '<input type="text" name="attention" class="dr-input"/>'
    if ri == 7 and col == 1:
        return "", '<input type="text" name="fax" class="dr-input"/>'
    if ri == 7 and col == 5:
        return "", '<input type="text" name="request_no" class="dr-input"/>'
    if ri == 8 and col == 5:
        return "", '<input type="text" name="manager" class="dr-input"/>'

    field_map = {11: "general_contractor", 12: "subcontractor", 13: "contract", 14: "contact_person"}
    if ri in field_map and col == 1:
        return "", f'<input type="text" name="{field_map[ri]}" class="dr-input"/>'

    if ri == 16 and col == 0:
        return "", '<input type="text" name="phone" class="dr-input"/>'
    if ri == 16 and col == 1:
        return "", '<input type="text" name="contractor_fax" class="dr-input"/>'
    if ri == 16 and col == 3:
        return "", '<input type="text" name="email" class="dr-input"/>'
    if ri == 16 and col == 5:
        return "", '<textarea name="extra_info" class="dr-textarea"></textarea>'

    if ri == 19 and col == 0:
        return "", '<textarea name="weather" class="dr-textarea"></textarea>'
    if ri == 19 and col == 1:
        inner = (
            '<label class="dr-check"><input type="checkbox" name="shift_day" checked/> '
            "08:00-20:00 (дневная смена)</label>"
            '<label class="dr-check"><input type="checkbox" name="shift_night"/> '
            "20:00-08:00 (ночная смена)</label>"
        )
        return "dr-checks", inner
    if ri == 19 and col == 4:
        return "", '<input type="text" name="prev_inspector" class="dr-input"/>'

    if ri in (22, 23, 24) and col == 0:
        n = ri - 21
        return "", f'<textarea name="sk_desc_{n}" class="dr-textarea"></textarea>'
    if ri in (22, 23, 24) and col == 4:
        n = ri - 21
        return "", f'<textarea name="sk_loc_{n}" class="dr-textarea"></textarea>'
    if ri in (22, 23, 24) and col == 5:
        n = ri - 21
        return "", f'<textarea name="sk_ref_{n}" class="dr-textarea"></textarea>'

    if ri in (27, 28) and col == 0:
        n = ri - 26
        return "", f'<textarea name="dup_desc_{n}" class="dr-textarea"></textarea>'
    if ri in (27, 28) and col == 4:
        n = ri - 26
        return "", f'<textarea name="dup_loc_{n}" class="dr-textarea"></textarea>'
    if ri in (27, 28) and col == 5:
        n = ri - 26
        return "", f'<textarea name="dup_act_{n}" class="dr-textarea"></textarea>'

    if ri == 30 and col == 0:
        return "", '<textarea name="remarks_photos_1" class="dr-textarea"></textarea>'
    if ri == 31 and col == 0:
        return "", '<textarea name="remarks_photos_2" class="dr-textarea"></textarea>'
    if ri == 33 and col == 0:
        return "", '<textarea name="remarks_peb_1" class="dr-textarea"></textarea>'
    if ri == 34 and col == 0:
        return "", '<textarea name="remarks_peb_2" class="dr-textarea"></textarea>'

    if ri == 35 and col == 2:
        return "", '<textarea name="recommendations" class="dr-textarea"></textarea>'
    if ri == 36 and col == 2:
        return "", '<textarea name="for_info" class="dr-textarea"></textarea>'
    if ri == 37 and col == 2:
        return "", '<textarea name="conclusion" class="dr-textarea"></textarea>'
    if ri == 39 and col == 0:
        return "", '<input type="text" name="engineer_name" class="dr-input"/>'

    if ri in (29, 32) and col == 0:
        return "dr-subsection", escape(text)

    if text:
        center_labels = {
            "Принято",
            "Обнаружено несоответствие",
            "Отложено",
            "Прочее",
            "подпись",
            "Описание действий",
            "Участок, ПК",
            "Ссылка",
            "№ акта, дата",
        }
        if text in center_labels or "Описание действий" in text:
            return "dr-label dr-center", escape(text).replace("\n", "<br>")
        return "dr-label", escape(text).replace("\n", "<br>")
    return "", "&nbsp;"


_STATUS_LABELS = (
    "Принято",
    "Обнаружено несоответствие",
    "Отложено",
    "Прочее",
)
_STATUS_VALUES = ("accepted", "nonconformity", "deferred", "other")


def _colgroup_html(layout: TableLayout) -> str:
    cols = "".join(f'<col class="dr-c{i}"/>' for i in range(len(layout.col_widths_dxa)))
    return f"<colgroup>{cols}</colgroup>"


def _tr_classes(ri: int) -> str:
    parts = [f"dr-r{ri}"]
    if ri in _ROW_SK:
        parts.append("dr-row-sk")
    elif ri in _ROW_TALL:
        parts.append("dr-row-tall")
    return " ".join(parts)


def _status_header_row() -> str:
    cells = [
        '<td class="dr-cell dr-section" colspan="2">СТАТУС СТРОИТЕЛЬНОГО КОНТРОЛЯ</td>',
    ]
    for label, value in zip(_STATUS_LABELS, _STATUS_VALUES, strict=True):
        cells.append(
            '<td class="dr-cell dr-status-head">'
            f'<input type="radio" name="sk_status" value="{value}"/>'
            f"<span>{escape(label)}</span></td>"
        )
    return '<tr class="dr-r17 dr-row-status">' + "".join(cells) + "</tr>"


def render_table_html(template_path: Path | None = None, layout: TableLayout | None = None) -> str:
    path = template_path or _TEMPLATE
    layout = layout or read_layout_from_docx(path)
    doc = Document(str(path))
    table = doc.tables[0]
    grid = _grid_rows(table)
    lines = ['<table class="dr-table">', _colgroup_html(layout), "<tbody>"]
    for ri, row in enumerate(grid):
        if ri == 17:
            lines.append(_status_header_row())
            continue
        lines.append(f'<tr class="{_tr_classes(ri)}">')
        for cell in row:
            if cell is None:
                lines.append('<td class="dr-cell">&nbsp;</td>')
                continue
            if cell == "SPAN":
                continue
            cs = cell["colspan"]
            cls, inner = _cell_content(ri, cell)
            class_attr = f"dr-cell {cls}".strip()
            span = f' colspan="{cs}"' if cs > 1 else ""
            lines.append(f'<td class="{class_attr}"{span}>{inner}</td>')
        lines.append("</tr>")
    lines.append("</tbody></table>")
    return "\n".join(lines)


def render_daily_report_page(template_path: Path | None = None) -> tuple[str, str]:
    """Возвращает (css из шаблона, html таблицы)."""
    path = template_path or _TEMPLATE
    layout = read_layout_from_docx(path)
    return render_layout_css(layout), render_table_html(path, layout)
