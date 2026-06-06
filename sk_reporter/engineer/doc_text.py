"""Извлечение текста из legacy .doc / .docx."""

from __future__ import annotations

import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path


def _read_txt_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _extract_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            line = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if line:
                parts.append(line)
    return "\n".join(parts)


def _extract_doc_mac(path: Path) -> str | None:
    if sys.platform != "darwin":
        return None
    textutil = shutil.which("textutil")
    if not textutil:
        return None
    proc = subprocess.run(
        [textutil, "-convert", "txt", "-stdout", str(path)],
        capture_output=True,
        text=True,
        errors="replace",
    )
    if proc.returncode == 0 and proc.stdout.strip():
        return proc.stdout
    return None


def _extract_doc_soffice(path: Path) -> str | None:
    for binary in ("soffice", "libreoffice"):
        exe = shutil.which(binary)
        if not exe:
            continue
        with tempfile.TemporaryDirectory() as tmp:
            out_dir = Path(tmp)
            proc = subprocess.run(
                [exe, "--headless", "--convert-to", "txt:Text", "--outdir", str(out_dir), str(path)],
                capture_output=True,
                text=True,
                errors="replace",
            )
            if proc.returncode != 0:
                continue
            txt_files = list(out_dir.glob("*.txt"))
            if txt_files:
                return _read_txt_file(txt_files[0])
    return None


def extract_doc_text(path: Path) -> str:
    path = Path(path)
    if not path.is_file():
        raise FileNotFoundError(path)

    if path.suffix.lower() == ".docx":
        return _extract_docx(path)

    text = _extract_doc_mac(path) or _extract_doc_soffice(path)
    if text:
        return text
    raise RuntimeError(
        f"Не удалось прочитать {path.name}: нужен macOS textutil или LibreOffice (soffice)"
    )


_SNIPPET_MARKERS = (
    "Контролируемые параметры",
    "Работы по",
    "Область применения",
    "Технический надзор",
)


def control_snippet_from_tk(text: str, max_chars: int = 900) -> str:
    """Выдержка из ТК для описания контроля в отчёте."""
    cleaned = re.sub(r"\s+", " ", text.replace("\x07", " ")).strip()
    if not cleaned:
        return ""

    for marker in _SNIPPET_MARKERS:
        idx = cleaned.find(marker)
        if idx >= 0:
            chunk = cleaned[idx : idx + max_chars * 2]
            chunk = re.sub(r" HYPERLINK[^\"]*\"[^\"]*\"", "", chunk)
            chunk = re.sub(r"\s+", " ", chunk).strip()
            if len(chunk) > max_chars:
                chunk = chunk[: max_chars - 1].rsplit(" ", 1)[0] + "…"
            return chunk

    chunk = cleaned[:max_chars]
    if len(cleaned) > max_chars:
        chunk = chunk.rsplit(" ", 1)[0] + "…"
    return chunk
