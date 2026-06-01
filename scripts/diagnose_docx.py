#!/usr/bin/env python3
"""Диагностика .docx отчёта: сетка таблицы, битые строки, картинки."""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from apply_template_layout import DEFAULT_GRID_COLS, diagnose_document, hardcoded_layout
from docx import Document


def main() -> None:
    if len(sys.argv) < 2:
        print("Использование: python3 scripts/diagnose_docx.py <файл.docx>")
        sys.exit(1)

    path = Path(sys.argv[1])
    if not path.is_file():
        print(f"Файл не найден: {path}")
        sys.exit(1)

    doc = Document(str(path))
    layout = hardcoded_layout()

    print(f"Файл: {path.name}")
    print(f"Таблиц: {len(doc.tables)}")
    for i, t in enumerate(doc.tables):
        print(f"  табл.{i + 1}: строк {len(t.rows)}, колонок (python-docx) {len(t.columns) if t.rows else 0}")

    warns = diagnose_document(doc, layout)
    if warns:
        print("ПРОБЛЕМЫ:")
        for w in warns:
            print(f"  - {w}")
    else:
        print("Сетка: все строки на 6 колонок (OK)")

    print(f"Инлайн-картинок: {len(doc.inline_shapes)}")
    print(f"Эталон колонок: {DEFAULT_GRID_COLS}")


if __name__ == "__main__":
    main()
