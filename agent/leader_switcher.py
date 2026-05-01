"""
Агент переключения руководителя в отчёте СК.
Меняет между Аниськовым В.И. и Манджиевым И.А.
"""

from docx import Document
from pathlib import Path
from typing import Literal

LEADERS = {
    "aniskov": {
        "row8_col4": "Руководитель",
        "row8_col5": "Аниськов Владимир Иванович",
        "row42_col4": "Руководитель проекта СК",
        "row42_col5": "Руководитель проекта СК",
        "row43_col4": "Аниськов Владимир Иванович",
        "row43_col5": "Аниськов Владимир Иванович",
    },
    "mandzhiev": {
        "row8_col4": "И.О. Руководителя",
        "row8_col5": "Манджиев Игорь Александрович",
        "row42_col4": "И.О. Руководителя проекта СК",
        "row42_col5": "И.О. Руководителя проекта СК",
        "row43_col4": "Манджиев Игорь Александрович",
        "row43_col5": "Манджиев Игорь Александрович",
    }
}


def switch_leader(filepath: str, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """
    Меняет руководителя в отчёте.
    
    Args:
        filepath: путь к .docx файлу
        leader: "aniskov" или "mandzhiev"
    
    Returns:
        (успех, сообщение)
    """
    try:
        doc = Document(filepath)
        
        if not doc.tables:
            return False, "В документе нет таблиц"
        
        table = doc.tables[0]
        data = LEADERS[leader]
        
        # Проверяем, что таблица достаточно большая
        if len(table.rows) < 44:
            return False, f"Таблица слишком мала: {len(table.rows)} строк (нужно минимум 44)"
        
        # Row 8: Колонки 4 и 5
        if len(table.rows[8].cells) >= 6:
            table.rows[8].cells[4].text = data["row8_col4"]
            table.rows[8].cells[5].text = data["row8_col5"]
        
        # Row 42: Колонки 4 и 5
        if len(table.rows[42].cells) >= 6:
            table.rows[42].cells[4].text = data["row42_col4"]
            table.rows[42].cells[5].text = data["row42_col5"]
        
        # Row 43: Колонки 4 и 5
        if len(table.rows[43].cells) >= 6:
            table.rows[43].cells[4].text = data["row43_col4"]
            table.rows[43].cells[5].text = data["row43_col5"]
        
        doc.save(filepath)
        
        leader_name = "Аниськов В.И." if leader == "aniskov" else "Манджиев И.А."
        return True, f"Установлен руководитель: {leader_name}"
        
    except Exception as e:
        return False, f"Ошибка: {str(e)}"


def detect_current_leader(filepath: str) -> str:
    """
    Определяет, какой руководитель сейчас в документе.
    
    Returns:
        "aniskov", "mandzhiev" или "unknown"
    """
    try:
        doc = Document(filepath)
        if not doc.tables:
            return "unknown"
        
        table = doc.tables[0]
        if len(table.rows) < 44:
            return "unknown"
        
        # Проверяем Row 8, Col 4
        if len(table.rows[8].cells) >= 5:
            text = table.rows[8].cells[4].text.strip().lower()
            if "манджиев" in text or "и.о." in text:
                return "mandzhiev"
            elif "аниськов" in text or ("руководитель" in text and "проекта" not in text):
                return "aniskov"
        
        # Проверяем Row 43, Col 4
        if len(table.rows[43].cells) >= 5:
            text = table.rows[43].cells[4].text.strip().lower()
            if "манджиев" in text:
                return "mandzhiev"
            elif "аниськов" in text:
                return "aniskov"
        
        return "unknown"
        
    except Exception:
        return "unknown"