"""
Парсер отчётов СК на базе Claude API (Anthropic).
Извлекает структурированные данные из ежедневных отчётов инженеров.
"""

import json
import os
import re
from pathlib import Path
from typing import Optional

import anthropic
from docx import Document

MODEL = "claude-haiku-4-5"

PARSE_PROMPT = """Ты — эксперт по разбору ежедневных отчётов строительного контроля (СК).

Извлеки из текста отчёта следующие поля и верни ТОЛЬКО валидный JSON без пояснений:

{
  "date": "дата отчёта в формате ДД.ММ.ГГГГ или null",
  "company": "название компании-подрядчика или null",
  "object": "название объекта/куста скважин или null",
  "inspector": "ФИО инспектора СК (строительного контролёра) или null",
  "weather": "погодные условия из отчёта или null",
  "work_done": "краткое описание выполненных работ (2-5 предложений) или null",
  "violations": "выявленные нарушения/замечания или null",
  "conclusion": "вывод/заключение инспектора или null",
  "report_number": "номер отчёта/куста или null",
  "work_volume": "объём выполненных работ с единицами измерения или null"
}

Правила:
- Верни ТОЛЬКО JSON, без markdown-блоков, без пояснений
- Если поле не найдено — ставь null
- Текст на русском — оставляй на русском
- Не придумывай данные, только то что есть в тексте
"""


def extract_full_text(filepath: str) -> str:
    """Извлекает весь текст из docx для отправки в Claude."""
    try:
        doc = Document(filepath)
        parts = []

        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                parts.append(t)

        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        row_texts.append(t)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        return "\n".join(parts)
    except Exception as e:
        print(f"[ERROR] extract_full_text {filepath}: {e}")
        return ""


def parse_report(filepath: str, api_key: Optional[str] = None) -> Optional[dict]:
    """
    Разбирает отчёт и возвращает структурированный dict.
    api_key берётся из параметра или переменной окружения ANTHROPIC_API_KEY.
    """
    key = api_key or os.environ.get("ANTHROPIC_API_KEY")
    if not key:
        print("[ERROR] ANTHROPIC_API_KEY не задан")
        return None

    text = extract_full_text(filepath)
    if not text:
        print(f"[ERROR] Не удалось извлечь текст из {filepath}")
        return None

    filename = Path(filepath).name
    print(f"[PARSE] Анализирую: {filename}")

    try:
        client = anthropic.Anthropic(api_key=key)
        response = client.messages.create(
            model=MODEL,
            max_tokens=1024,
            system=PARSE_PROMPT,
            messages=[
                {
                    "role": "user",
                    "content": f"Разбери отчёт:\n\nИмя файла: {filename}\n\n{text[:4000]}"
                }
            ],
        )

        raw = response.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)

        data = json.loads(raw)
        data["_source_file"] = filename
        print(f"[PARSE] OK: {filename} → {data.get('company')} / {data.get('date')}")
        return data

    except json.JSONDecodeError as e:
        print(f"[ERROR] JSON parse failed for {filename}: {e}")
        return None
    except Exception as e:
        print(f"[ERROR] Claude API error for {filename}: {e}")
        return None


def parse_reports_batch(filepaths: list[str], api_key: Optional[str] = None) -> list[dict]:
    """Разбирает список отчётов и возвращает список результатов."""
    results = []
    for fp in filepaths:
        result = parse_report(fp, api_key)
        if result:
            results.append(result)
        else:
            results.append({
                "_source_file": Path(fp).name,
                "_parse_error": True,
            })
    return results
