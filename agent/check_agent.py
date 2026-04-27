"""
Агент проверки загруженных отчётов строительного контроля.
Проверяет объемы работ, удаляет работы с 0 суточным объемом,
проверяет соответствие описаний действий типам работ.
Возвращает отчет об ошибках и рекомендациях (без изменения файлов).
"""

import re
import json
from typing import Optional, List, Dict, Any
from pathlib import Path
from docx import Document

MODEL = "gemma4:31b-cloud"

SYSTEM_PROMPT = """Ты — ведущий инженер строительного контроля, проверяешь качество заполнения ежедневных отчётов.

Правила проверки:
1. ОБЪЕМЫ: накопительный объем НЕ МОЖЕТ быть больше проектного объема. Это критическая ошибка!
2. НУЛЕВЫЕ РАБОТЫ: если суточный объем = 0, эту работу нужно УДАЛИТЬ из отчета. Это ошибка!
3. ОПИСАНИЯ: описание работы должно точно соответствовать типу выполняемой работы.

Для КАЖДОЙ проблемы напиши:
- В чем суть проблемы (с конкретными цифрами и названиями)
- Почему это проблема
- Как исправить

Будь жесток - это строительство, ошибки дорогие!
Используй четкие примеры из отчета."""


def extract_report_data(filepath: str) -> Dict[str, Any]:
    """
    Извлекает данные о работах из DOCX отчета.
    Возвращает структурированные данные о работах с объемами.
    """
    try:
        doc = Document(filepath)
        works = []
        work_counter = 0

        # Ищем в таблицах и тексте работы с объемами
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " ".join(cells)

                # Пытаемся найти паттерн: РАБОТА: 1. ... Проектный объем – X ... Объем за сутки – Y ...
                if "Проектный объем" in row_text or "проектный объем" in row_text.lower():
                    work_counter += 1
                    # Собираем данные о работе
                    work_data = {
                        "id": work_counter,
                        "description": cells[0] if cells else "",
                        "project_volume": None,
                        "daily_volume": None,
                        "cumulative_volume": None,
                    }

                    # Парсим объемы
                    project_match = re.search(r"Проектный объем[:\s–]+([0-9.,]+)", row_text)
                    daily_match = re.search(r"Объем за сутки[:\s–]+([0-9.,]+)", row_text)
                    cumul_match = re.search(r"Накопительный объем[:\s–]+([0-9.,]+)", row_text)

                    if project_match:
                        work_data["project_volume"] = float(project_match.group(1).replace(",", "."))
                    if daily_match:
                        work_data["daily_volume"] = float(daily_match.group(1).replace(",", "."))
                    if cumul_match:
                        work_data["cumulative_volume"] = float(cumul_match.group(1).replace(",", "."))

                    works.append(work_data)

        # Если в таблицах не нашли, парсим текст параграфов
        if not works:
            text_parts = []
            for para in doc.paragraphs:
                text_parts.append(para.text)
            full_text = "\n".join(text_parts)

            # Ищем все работы в формате "РАБОТА: N. ..."
            work_sections = re.split(r"РАБОТА:\s*\d+\.", full_text)
            for i, section in enumerate(work_sections[1:], 1):
                work_data = {
                    "id": i,
                    "description": section[:100],  # первые 100 символов
                    "project_volume": None,
                    "daily_volume": None,
                    "cumulative_volume": None,
                }

                project_match = re.search(r"Проектный объем[:\s–]+([0-9.,]+)", section)
                daily_match = re.search(r"Объем за сутки[:\s–]+([0-9.,]+)", section)
                cumul_match = re.search(r"Накопительный объем[:\s–]+([0-9.,]+)", section)

                if project_match:
                    work_data["project_volume"] = float(project_match.group(1).replace(",", "."))
                if daily_match:
                    work_data["daily_volume"] = float(daily_match.group(1).replace(",", "."))
                if cumul_match:
                    work_data["cumulative_volume"] = float(cumul_match.group(1).replace(",", "."))

                works.append(work_data)

        return {
            "filepath": filepath,
            "filename": Path(filepath).name,
            "works": works,
            "_source_file": Path(filepath).name,
        }
    except Exception as e:
        print(f"[CHECK_AGENT] extract_report_data error: {e}")
        return {
            "filepath": filepath,
            "filename": Path(filepath).name,
            "works": [],
            "error": str(e),
            "_source_file": Path(filepath).name,
        }


def check_report(filepath: str) -> dict:
    """
    Основная функция агента.
    Читает отчет, проверяет его и возвращает результаты.
    """
    # Извлекаем данные
    report_data = extract_report_data(filepath)
    works = report_data.get("works", [])

    if not works:
        return {
            "ok": False,
            "errors": [{"type": "parse_error", "message": "Не удалось распарсить работы из отчета"}],
            "recommendations": ["Проверьте формат документа"],
            "summary": "Ошибка парсинга",
            "_source_file": report_data.get("_source_file"),
        }

    # Подготавливаем промпт для агента
    works_text = json.dumps(works, ensure_ascii=False, indent=2)

    # Создаем промпт для LLM
    user_prompt = f"""Проверь отчет. Вот данные работ:

{works_text}

НАПИШИ подробно:
1. Все найденные ошибки с цифрами и названиями
2. Почему это ошибки
3. Как их исправить

Конкретные примеры! Не общие слова!"""

    # Вызываем ollama
    try:
        import ollama
        response = ollama.chat(
            model=MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt}
            ],
            stream=False,
        )

        report_text = response.get("message", {}).get("content", "").strip()

        if not report_text:
            return {
                "ok": False,
                "report": "Ошибка: пустой ответ модели",
                "_source_file": report_data.get("_source_file"),
            }

        # Проверяем есть ли ошибки по ключевым словам
        has_errors = any(word in report_text.lower() for word in ["ошибка", "проблема", "некорректно", "неправильно", "превышает", "удалить"])

        result = {
            "ok": not has_errors,
            "report": report_text,
            "_source_file": report_data.get("_source_file"),
        }

        print(f"[CHECK_AGENT] {'OK' if result['ok'] else 'ERRORS'}: {report_data.get('filename')}")
        return result

    except Exception as e:
        print(f"[CHECK_AGENT] Error calling ollama: {e}")
        return {
            "ok": False,
            "errors": [{"type": "model_error", "message": str(e)}],
            "recommendations": ["Попробуйте позже"],
            "summary": f"Ошибка модели: {e}",
            "_source_file": report_data.get("_source_file"),
        }
