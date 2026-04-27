"""
Агент на базе Ollama + qwen3.5:cloud для анализа и сборки отчётов СК.
Определяет компанию по содержимому документа, вставляет в болванку.
"""

import ollama
import re
import json
import zipfile
import shutil
import os
from copy import deepcopy
from pathlib import Path
from typing import Optional
from lxml import etree

from docx import Document

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))
from companies import COMPANIES

# Преобразуем в формат для агента: название → список ключевых слов
COMPANIES_MAP = {name: keywords for name, keywords in COMPANIES}
KNOWN_COMPANIES = [name for name, _ in COMPANIES]

MODEL = "qwen3.5:cloud"

SYSTEM_PROMPT = f"""Ты определяешь компанию-подрядчика по тексту отчёта строительного контроля.

Список компаний: {json.dumps(KNOWN_COMPANIES, ensure_ascii=False)}

Правила:
1. Ищи название компании в шапке, подписях, таблицах документа.
2. Верни ТОЛЬКО название из списка выше — без пояснений, без кавычек.
3. Если не нашёл — верни UNKNOWN.
"""


def extract_text(filepath: str) -> str:
    try:
        doc = Document(filepath)
        parts = []
        for para in doc.paragraphs[:50]:
            t = para.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables[:3]:
            for row in table.rows[:5]:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and len(t) > 2:
                        parts.append(t)
        return "\n".join(parts)
    except Exception as e:
        print(f"[ERROR] extract_text {filepath}: {e}")
        return ""


def detect_geodesy(filepath: str) -> Optional[str]:
    """
    Проверяет, является ли отчёт геодезическим контролем.
    Ищет в таблице ячейку [2,3] (Направление контроля).
    """
    try:
        doc = Document(filepath)
        if len(doc.tables) == 0:
            return None

        table = doc.tables[0]
        if len(table.rows) < 3:
            return None

        # Проверяем ячейку [2,3] — "Направление контроля"
        row = table.rows[2]
        if len(row.cells) < 4:
            return None

        direction_cell = row.cells[3].text.strip().lower()

        if "геодезический" in direction_cell:
            filename = Path(filepath).name
            print(f"[GEODESY] {filename} → Геодезический контроль")
            return "Геодезический контроль"

        return None
    except Exception as e:
        print(f"[ERROR] detect_geodesy {filepath}: {e}")
        return None


def detect_company(filepath: str) -> Optional[str]:
    filename = Path(filepath).name
    filename_lower = filename.lower()

    # Шаг 0: проверка геодезии (специальный случай)
    geodesy = detect_geodesy(filepath)
    if geodesy:
        return geodesy

    # Шаг 1: проверка по всем вариантам написания в имени файла
    for company, keywords in COMPANIES_MAP.items():
        if any(kw in filename_lower for kw in keywords):
            print(f"[FILENAME] {filename} → {company}")
            return company

    # Шаг 2: AI анализирует содержимое документа
    text = extract_text(filepath)
    if not text:
        print(f"[UNKNOWN] Не удалось прочитать текст: {filename}")
        return None

    try:
        response = ollama.chat(
            model=MODEL,
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": f"Определи компанию:\n\n{text[:2000]}"},
            ],
            options={"temperature": 0.0, "num_predict": 30},
        )
        answer = response["message"]["content"].strip()
        answer = re.sub(r"<think>.*?</think>", "", answer, flags=re.DOTALL).strip()
        clean = re.sub(r'["\'\.\n]', "", answer).strip()

        for company in KNOWN_COMPANIES:
            if company.lower() in clean.lower() or clean.lower() in company.lower():
                print(f"[AI] {filename} → {company}")
                return company

        print(f"[UNKNOWN] AI не распознал компанию для: {filename}")
        return None

    except Exception as e:
        print(f"[ERROR] ollama: {e}")
        return None


def _xml_bytes(root) -> bytes:
    """Сериализует XML с двойными кавычками в заголовке — Word на Windows требует этого."""
    body = etree.tostring(root, encoding="unicode")
    return ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + body).encode("utf-8")


def merge_report_into_template(template_path: str, report_path: str, output_path: str) -> bool:
    """
    Вставляет содержимое report_path в конец template_path через ZIP.
    Картинки копируются с переименованием чтобы избежать дублей.
    """
    try:
        import tempfile

        # Читаем XML тела отчёта
        with zipfile.ZipFile(report_path, "r") as zr:
            doc_xml = zr.read("word/document.xml")
            report_media = {
                name: zr.read(name)
                for name in zr.namelist()
                if name.startswith("word/media/")
            }
            try:
                rels_xml = zr.read("word/_rels/document.xml.rels")
            except KeyError:
                rels_xml = None

        # Разбираем XML отчёта
        src_root = etree.fromstring(doc_xml)
        NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        body = src_root.find(f"{{{NS}}}body")

        # Собираем элементы тела (без последнего sectPr)
        body_elements = []
        for el in body:
            tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            if tag != "sectPr":
                body_elements.append(deepcopy(el))

        # Читаем rels отчёта чтобы знать rId→имя файла
        rid_to_name: dict[str, str] = {}
        if rels_xml:
            rels_root = etree.fromstring(rels_xml)
            for rel in rels_root:
                rId = rel.get("Id", "")
                target = rel.get("Target", "")
                if target.startswith("media/"):
                    rid_to_name[rId] = "word/" + target

        # Копируем шаблон во временный файл, потом в output
        tmp = output_path + ".build.docx"
        shutil.copy2(template_path, tmp)

        # Определяем какие медиафайлы уже есть в шаблоне
        with zipfile.ZipFile(tmp, "r") as zt:
            existing_media = set(n for n in zt.namelist() if n.startswith("word/media/"))
            content_types_xml = zt.read("[Content_Types].xml")

        # Генерируем уникальные имена для медиафайлов отчёта
        name_map: dict[str, str] = {}  # старое имя → новое имя
        counter = len(existing_media) + 1
        for old_name, data in report_media.items():
            ext = Path(old_name).suffix
            new_name = f"word/media/img_r{counter}{ext}"
            while new_name in existing_media or new_name in name_map.values():
                counter += 1
                new_name = f"word/media/img_r{counter}{ext}"
            name_map[old_name] = new_name
            counter += 1

        # Читаем rels шаблона чтобы знать следующий свободный rId
        with zipfile.ZipFile(tmp, "r") as zt:
            master_rels_xml = zt.read("word/_rels/document.xml.rels")

        master_rels_root = etree.fromstring(master_rels_xml)
        existing_ids = set()
        for r in master_rels_root:
            rid = r.get("Id", "")
            if rid.startswith("rId"):
                try:
                    existing_ids.add(int(rid[3:]))
                except ValueError:
                    pass
        next_id = max(existing_ids, default=0) + 1

        # Строим финальный rid_map: старый rId отчёта → новый rId в шаблоне
        IMG_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
        rid_final_map: dict[str, str] = {}  # старый rId → новый rId
        new_rels: list[tuple[str, str]] = []  # (новый rId, target относительно word/)
        for old_rid, old_name in rid_to_name.items():
            if old_name in name_map:
                new_rid = f"rId{next_id}"
                next_id += 1
                new_target = "media/" + Path(name_map[old_name]).name
                rid_final_map[old_rid] = new_rid
                new_rels.append((new_rid, new_target))

        # Строим новый ZIP
        with zipfile.ZipFile(tmp, "r") as zt:
            with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zo:
                for item in zt.namelist():
                    if item == "word/document.xml":
                        master_xml = zt.read(item)
                        master_root = etree.fromstring(master_xml)
                        master_body = master_root.find(f"{{{NS}}}body")
                        last = master_body[-1] if len(master_body) else None
                        last_tag = last.tag.split("}")[-1] if last is not None else ""
                        if last_tag == "sectPr":
                            sect_pr = deepcopy(last)
                            master_body.remove(last)
                        else:
                            sect_pr = None
                        for el in body_elements:
                            el_copy = deepcopy(el)
                            if rid_final_map:
                                _patch_rids_by_target(el_copy, rid_final_map)
                            master_body.append(el_copy)
                        if sect_pr is not None:
                            master_body.append(sect_pr)
                        zo.writestr(item, _xml_bytes(master_root))
                    elif item == "word/_rels/document.xml.rels":
                        rels_data = zt.read(item)
                        rels_root = etree.fromstring(rels_data)
                        for new_rid, new_target in new_rels:
                            etree.SubElement(rels_root, "Relationship", {
                                "Id": new_rid,
                                "Type": IMG_REL,
                                "Target": new_target,
                            })
                        zo.writestr(item, _xml_bytes(rels_root))
                    elif item == "[Content_Types].xml":
                        ct_root = etree.fromstring(content_types_xml)
                        CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
                        existing_parts = set(el.get("PartName", "") for el in ct_root)
                        EXT_MAP = {
                            ".png": "image/png", ".jpg": "image/jpeg",
                            ".jpeg": "image/jpeg", ".gif": "image/gif",
                            ".bmp": "image/bmp", ".tiff": "image/tiff",
                        }
                        for new_name in name_map.values():
                            part = "/" + new_name
                            if part not in existing_parts:
                                ext = Path(new_name).suffix.lower()
                                ct = EXT_MAP.get(ext, "image/png")
                                etree.SubElement(ct_root, f"{{{CT_NS}}}Override", {
                                    "PartName": part,
                                    "ContentType": ct,
                                })
                        zo.writestr(item, _xml_bytes(ct_root))
                    else:
                        zo.writestr(item, zt.read(item))

                # Записываем новые медиафайлы
                for old_name, new_name in name_map.items():
                    if old_name in report_media:
                        zo.writestr(new_name, report_media[old_name])

        os.remove(tmp)
        return True

    except Exception as e:
        import traceback
        print(f"[ERROR] merge {report_path}: {e}")
        traceback.print_exc()
        return False


def _patch_rids_by_target(element, rid_new_map: dict) -> None:
    """Патчит r:embed/r:id атрибуты в XML элементе."""
    ATTRS = (
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed",
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id",
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link",
    )
    for attr in ATTRS:
        val = element.get(attr)
        if val and val in rid_new_map:
            element.set(attr, rid_new_map[val])
    for child in element:
        _patch_rids_by_target(child, rid_new_map)


