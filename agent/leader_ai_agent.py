"""AI Агент переключения руководителя через Ollama.
Использует LLM для анализа и замены.
"""

import json
import re
import ollama
from docx import Document
from pathlib import Path
from typing import Literal

MODEL = "qwen3.5:cloud"


def _ask_llm_for_leader_cells(cells_text: list) -> dict:
    """Ask LLM to find cells containing leader FIO and title/position."""
    cell_list = "\n".join(
        f"[{ti},{ri},{ci}]: {preview[:120]!r}"
        for ti, ri, ci, preview in cells_text
    )

    prompt = f"""В документе отчёта строительного контроля есть ячейки таблицы.
Найди ячейки, содержащие:
1. ФИО руководителя (фамилия имя отчество, например "Аниськов Владимир Иванович" или "Манджиев Игорь Александрович" — может быть написано с опечатками)
2. Должность/заголовок руководителя (например "Руководитель", "И.О. Руководителя", "Руководитель проекта СК", "И.О. Руководителя проекта СК" — может быть написано с опечатками или в нижнем регистре)

Список ячеек:
{cell_list}

Ответь ТОЛЬКО JSON без пояснений:
{{"fio_cell": [таблица, строка, столбец], "title_cell": [таблица, строка, столбец], "project_cell": [таблица, строка, столбец]}}

Примечания:
- fio_cell — ячейка с ФИО руководителя (не инженера!)
- title_cell — ячейка с заголовком ("Руководитель" / "И.О. Руководителя") — краткий вариант без "проекта СК"
- project_cell — ячейка с должностью ("Руководитель проекта СК" / "И.О. Руководителя проекта СК")
- Если ячейка не найдена — используй null
- title_cell и project_cell могут быть одной и той же ячейкой или разными
- НЕ путай инженера СК с руководителем проекта СК"""

    response = ollama.chat(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        options={"temperature": 0.0},
        stream=False,
    )
    raw = response.get("message", {}).get("content", "").strip()
    print(f"[LEADER_AGENT] LLM cell response: {raw[:400]}")

    json_match = re.search(r"\{[^{}]*\}", raw, re.DOTALL)
    if not json_match:
        return {}
    try:
        return json.loads(json_match.group())
    except Exception:
        return {}


def _force_write_para(cell, new_text: str):
    """Overwrite first paragraph of cell with new_text, preserve run formatting."""
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def _switch_single_file(filepath: str, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """Обрабатывает один файл — LLM находит нужные ячейки, принудительно записывает значения."""
    try:
        doc = Document(filepath)

        if not doc.tables:
            return False, "Нет таблиц в документе"

        if leader == "aniskov":
            target_fio = "Аниськов Владимир Иванович"
            target_title = "Руководитель"
            target_project = "Руководитель проекта СК"
        else:
            target_fio = "Манджиев Игорь Александрович"
            target_title = "И.О. Руководителя"
            target_project = "И.О. Руководителя проекта СК"

        cells = []
        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    txt = cell.text.strip()
                    if txt:
                        cells.append((ti, ri, ci, txt[:300]))

        coords = _ask_llm_for_leader_cells(cells)
        print(f"[LEADER_AGENT] LLM identified cells: {coords}")

        written = []

        def _write_coord(key, value):
            coord = coords.get(key)
            if coord and isinstance(coord, list) and len(coord) == 3:
                ti, ri, ci = coord
                try:
                    target_cell = doc.tables[ti].rows[ri].cells[ci]
                    _force_write_para(target_cell, value)
                    written.append(f"{key}→[{ti},{ri},{ci}]")
                except Exception as e:
                    print(f"[LEADER_AGENT] Failed to write {key}: {e}")

        _write_coord("fio_cell", target_fio)
        _write_coord("title_cell", target_title)
        _write_coord("project_cell", target_project)

        if not written:
            return False, "LLM не нашёл ячейки руководителя в документе"

        doc.save(filepath)
        filename = Path(filepath).name
        return True, f"→ {filename}: записано {', '.join(written)}"

    except Exception as e:
        filename = Path(filepath).name
        return False, f"→ {filename}: ошибка - {str(e)}"


def switch_leader_ai(filepaths: list, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """Обрабатывает список файлов."""
    if not filepaths:
        return False, "Нет файлов для обработки"

    results = []
    success_count = 0

    for filepath in filepaths:
        ok, msg = _switch_single_file(filepath, leader)
        results.append(msg)
        if ok:
            success_count += 1

    if success_count == 0:
        return False, "Ни один файл не обработан: " + "; ".join(results)

    output = "\n".join(results)
    output += f"\nОбработано: {success_count}/{len(filepaths)} файлов"
    return True, output


def switch_leader(filepath: str, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """Переключает руководителя в одном файле (совместимость со старым API)."""
    return _switch_single_file(filepath, leader)


_LEADER_TITLE_RE = re.compile(
    r"^(и\.?\s*о\.?\s*руководител[ья]|руководител[ья])(\s+проекта\s+ск)?$",
    re.IGNORECASE,
)
_PROJECT_RE = re.compile(r"проекта\s+ск", re.IGNORECASE)
_ENGINEER_RE = re.compile(r"инженер", re.IGNORECASE)


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def _force_write_para(cell, new_text: str):
    """Overwrite first paragraph of cell with new_text, preserve run formatting."""
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


def _switch_single_file(filepath: str, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """Быстрая замена руководителя по ключевым словам, без LLM."""
    try:
        doc = Document(filepath)

        if not doc.tables:
            return False, "Нет таблиц в документе"

        if leader == "aniskov":
            target_fio = "Аниськов Владимир Иванович"
            target_title = "Руководитель"
            target_project = "Руководитель проекта СК"
        else:
            target_fio = "Манджиев Игорь Александрович"
            target_title = "И.О. Руководителя"
            target_project = "И.О. Руководителя проекта СК"

        written = []

        for ti, table in enumerate(doc.tables):
            for ri, row in enumerate(table.rows):
                cells = row.cells
                texts = [_norm(c.text) for c in cells]

                # Check if this row contains any engineer label — skip entirely
                if any(_ENGINEER_RE.search(t) for t in texts):
                    continue

                # Find unique cell objects (merged cells repeat the same object)
                seen_ids = set()
                unique_cells = []
                for ci, cell in enumerate(cells):
                    cid = id(cell)
                    if cid not in seen_ids:
                        seen_ids.add(cid)
                        unique_cells.append((ci, cell, texts[ci]))

                for ci, cell, txt in unique_cells:
                    if not _LEADER_TITLE_RE.match(txt):
                        continue

                    # This cell is a leader label — determine which kind
                    if _PROJECT_RE.search(txt):
                        _force_write_para(cell, target_project)
                        written.append(f"project→[{ti},{ri},{ci}]")
                    else:
                        _force_write_para(cell, target_title)
                        written.append(f"title→[{ti},{ri},{ci}]")

                    # FIO may be in the same row (next unique cell) or the next row at same column
                    fio_written = False

                    # Check same row: next unique cell after the label
                    next_in_row = [
                        (nci, nc, ntxt)
                        for nci, nc, ntxt in unique_cells
                        if nci > ci and not _LEADER_TITLE_RE.match(ntxt)
                           and not _ENGINEER_RE.search(ntxt)
                    ]
                    if next_in_row:
                        nci, nc, _ = next_in_row[0]
                        _force_write_para(nc, target_fio)
                        written.append(f"fio→[{ti},{ri},{nci}]")
                        fio_written = True

                    # Also check next row at same column (for project section rows)
                    if ri + 1 < len(table.rows):
                        next_row_cells = table.rows[ri + 1].cells
                        next_row_texts = [_norm(c.text) for c in next_row_cells]
                        if not any(_ENGINEER_RE.search(t) for t in next_row_texts):
                            seen_next = set()
                            for nci, nc in enumerate(next_row_cells):
                                nid = id(nc)
                                if nid in seen_next:
                                    continue
                                seen_next.add(nid)
                                if nci == ci and not _LEADER_TITLE_RE.match(_norm(nc.text)):
                                    _force_write_para(nc, target_fio)
                                    written.append(f"fio→[{ti},{ri+1},{nci}]")
                                    break

        if not written:
            return False, "Ячейки руководителя не найдены в документе"

        doc.save(filepath)
        filename = Path(filepath).name
        return True, f"→ {filename}: {', '.join(written)}"

    except Exception as e:
        filename = Path(filepath).name
        return False, f"→ {filename}: ошибка - {str(e)}"


def switch_leader_ai(filepaths: list, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """Обрабатывает список файлов."""
    if not filepaths:
        return False, "Нет файлов для обработки"
    
    results = []
    success_count = 0

    for filepath in filepaths:
        ok, msg = _switch_single_file(filepath, leader)
        results.append(msg)
        if ok:
            success_count += 1

    if success_count == 0:
        return False, "Ни один файл не обработан: " + "; ".join(results)

    output = "\n".join(results)
    output += f"\nОбработано: {success_count}/{len(filepaths)} файлов"
    return True, output


def switch_leader(filepath: str, leader: Literal["aniskov", "mandzhiev"]) -> tuple[bool, str]:
    """Переключает руководителя в одном файле (совместимость со старым API)."""
    return _switch_single_file(filepath, leader)



