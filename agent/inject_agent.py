"""
Агент инъекции: берёт оригинальный docx + исправленный текст от check_agent,
использует LLM для нахождения нужных ячеек таблицы и вставляет исправленный текст.
"""

import json
import re
import shutil
import tempfile
from pathlib import Path
from docx import Document

MODEL = "gemma4:31b-cloud"


def _extract_cells(doc: Document) -> list:
    """Returns list of (table_idx, row_idx, col_idx, text_preview)."""
    cells = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if txt:
                    cells.append((ti, ri, ci, txt[:300]))
    return cells


def _ask_llm_for_cells(cells: list, corrected_text: str) -> dict:
    """Ask LLM to identify which cell indices contain part1 and part2 content."""
    import ollama

    cell_list = "\n".join(
        f"[{ti},{ri},{ci}]: {preview[:150]!r}"
        for ti, ri, ci, preview in cells
    )

    # Extract previews for context — works with or without ЧАСТЬ headers
    p2_start = re.search(r"(Наряд.допуск|Работы ведутся)", corrected_text, re.IGNORECASE)
    if p2_start:
        part1_preview = corrected_text[:p2_start.start()].strip()[-200:]
        part2_preview = corrected_text[p2_start.start():p2_start.start()+200]
    else:
        part1_preview = corrected_text[:200]
        part2_preview = ""

    prompt = f"""Ниже — список ячеек таблицы из документа docx. Каждая запись: [таблица,строка,столбец]: "начало текста ячейки".

{cell_list}

Тебе нужно найти:
1. Ячейку, содержащую ЧАСТЬ 1 отчёта — список видов работ с объёмами (начинается примерно с «Инспекционный контроль по» или содержит «Проектный объем», «Объем за сутки»). Начало ожидаемого нового содержимого: {part1_preview[:100]!r}
2. Ячейку, содержащую ЧАСТЬ 2 отчёта — описания выполненных работ (начинается примерно с «Наряд-допуск» или «Работы ведутся»). Начало ожидаемого нового содержимого: {part2_preview[:100]!r}

Ответь ТОЛЬКО JSON без пояснений:
{{"part1": [таблица, строка, столбец], "part2": [таблица, строка, столбец]}}

Если ячейка не найдена — используй null вместо массива."""

    response = ollama.chat(
        model=MODEL,
        messages=[{"role": "user", "content": prompt}],
        stream=False,
    )
    raw = response.get("message", {}).get("content", "").strip()
    print(f"[INJECT_AGENT] LLM cell response: {raw[:300]}")

    # Extract JSON from response
    json_match = re.search(r"\{[^{}]*\}", raw, re.DOTALL)
    if not json_match:
        return {}
    return json.loads(json_match.group())


def _parse_parts(corrected_text: str):
    """Parse part1 and part2 from LLM output.

    LLM may or may not output ЧАСТЬ 1 / ЧАСТЬ 2 headers.
    Part 1 starts at 'Инспекционный контроль' (or first numbered item with volumes).
    Part 2 starts at 'Наряд-допуск' or 'Работы ведутся'.
    """
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", corrected_text)

    # Get text after ## ИСПРАВЛЕННЫЙ ОТЧЁТ if present
    section_match = re.search(
        r"##\s*ИСПРАВЛЕННЫЙ\s*ОТЧЁТ[^\n]*\n(.*?)$", cleaned, re.DOTALL | re.IGNORECASE
    )
    search_text = section_match.group(1).strip() if section_match else cleaned.strip()

    # Try explicit ЧАСТЬ markers first
    part1_match = re.search(
        r"ЧАСТЬ\s*1[^:\n]*[:\n](.*?)(?=ЧАСТЬ\s*2\b|$)", search_text, re.DOTALL | re.IGNORECASE
    )
    part2_match = re.search(
        r"ЧАСТЬ\s*2[^:\n]*[:\n](.*?)$", search_text, re.DOTALL | re.IGNORECASE
    )

    # Fallback: split by part2 start marker
    if not part1_match or not part2_match:
        p2_start = re.search(
            r"(Наряд.допуск|Работы ведутся)", search_text, re.IGNORECASE
        )
        if p2_start:
            raw_part1 = search_text[:p2_start.start()].strip()
            raw_part2 = search_text[p2_start.start():].strip()
            # Strip any ЧАСТЬ 1 header line from part1
            raw_part1 = re.sub(r"^ЧАСТЬ\s*1[^\n]*\n", "", raw_part1, flags=re.IGNORECASE).strip()
            part1_lines = [l.rstrip() for l in raw_part1.splitlines()] if raw_part1 else []
            part2_lines = [l.rstrip() for l in raw_part2.splitlines()] if raw_part2 else []
            print(f"[INJECT_AGENT] fallback parse: part1={len(part1_lines)} lines, part2={len(part2_lines)} lines")
            return part1_lines, part2_lines

    part1_lines = []
    part2_lines = []
    if part1_match:
        part1_lines = [l.rstrip() for l in part1_match.group(1).strip().splitlines()]
    if part2_match:
        part2_lines = [l.rstrip() for l in part2_match.group(1).strip().splitlines()]

    print(f"[INJECT_AGENT] parsed part1={len(part1_lines)} lines, part2={len(part2_lines)} lines")
    return part1_lines, part2_lines


def _write_parts_to_cell(cell, part1_lines: list, part2_lines: list):
    """Insert part1 and part2 at the beginning of cell, before original text."""
    # Combine both parts
    all_lines = []
    if part1_lines:
        all_lines.extend(part1_lines)
    if part2_lines:
        all_lines.extend(part2_lines)
    
    if not all_lines:
        return
    
    # Get the first paragraph (header "Описание действий")
    insert_point = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    
    # Insert all lines after the header (in reverse to keep order)
    for line in reversed(all_lines):
        new_para = cell.add_paragraph(line)
        insert_point._element.addnext(new_para._element)


def inject_into_docx(filepath: str, corrected_text: str, source_filename: str) -> dict:
    stem = Path(source_filename).stem
    try:
        print(f"[INJECT_AGENT] === FULL corrected_text ===\n{corrected_text[:500]}\n... (truncated) ===")
        part1_lines, part2_lines = _parse_parts(corrected_text)

        if not part1_lines and not part2_lines:
            return {
                "ok": False,
                "error": "Не удалось распарсить ЧАСТЬ 1 / ЧАСТЬ 2 из ответа агента",
                "docx_path": None,
            }

        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir) / Path(filepath).name
            shutil.copy2(filepath, tmp_path)

            doc = Document(str(tmp_path))
            
            # Find cell with "Описание действий" header
            target_cell = None
            for ti, table in enumerate(doc.tables):
                for ri, row in enumerate(table.rows):
                    for ci, cell in enumerate(row.cells):
                        if "Описание действий" in cell.text:
                            target_cell = cell
                            print(f"[INJECT_AGENT] Found 'Описание действий' at [{ti},{ri},{ci}]")
                            break
                    if target_cell:
                        break
                if target_cell:
                    break
            
            if target_cell and (part1_lines or part2_lines):
                # Find table and row index
                target_table = None
                target_row_idx = None
                for ti, table in enumerate(doc.tables):
                    for ri, row in enumerate(table.rows):
                        for ci, cell in enumerate(row.cells):
                            if cell == target_cell:
                                target_table = table
                                target_row_idx = ri
                                break
                        if target_row_idx is not None:
                            break
                    if target_row_idx is not None:
                        break
                
                if target_table is not None:
                    # Add two new rows
                    col_count = len(target_table.rows[0].cells) if target_table.rows else 0
                    
                    # Add row for Part 1
                    new_row_1 = target_table.add_row()
                    if part1_lines and col_count > 0:
                        new_row_1.cells[0].text = "\n".join(part1_lines)
                    
                    # Add row for Part 2  
                    new_row_2 = target_table.add_row()
                    if part2_lines and col_count > 0:
                        new_row_2.cells[0].text = "\n".join(part2_lines)
                    
                    # Move new rows to correct position (after target_row_idx)
                    tbl = target_table._tbl
                    tr_elements = list(tbl.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr'))
                    
                    if len(tr_elements) >= 2 and target_row_idx < len(tr_elements) - 2:
                        # Get the rows we just added (last two)
                        tr_1 = tr_elements[-2]
                        tr_2 = tr_elements[-1]
                        ref_tr = tr_elements[target_row_idx]
                        
                        # Remove from end
                        tbl.remove(tr_2)
                        tbl.remove(tr_1)
                        
                        # Insert after reference row
                        ref_tr.addnext(tr_2)
                        ref_tr.addnext(tr_1)
                    
                    print(f"[INJECT_AGENT] Inserted 2 new rows with Part 1 and Part 2")

            output_dir = Path(__file__).parent.parent / "output"
            output_dir.mkdir(exist_ok=True)
            final_path = output_dir / f"{stem}_исправлен.docx"
            doc.save(str(final_path))

        print(f"[INJECT_AGENT] Saved: {final_path}")
        return {"ok": True, "docx_path": str(final_path)}

    except Exception as e:
        import traceback
        print(f"[INJECT_AGENT] error: {e}\n{traceback.format_exc()}")
        return {"ok": False, "error": str(e), "docx_path": None}
