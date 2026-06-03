#!/usr/bin/env python3
"""Анализ шаблона и отчётов — выводит размеры таблиц, строк и ячеек."""

from docx import Document
from docx.oxml.ns import qn
import sys
from pathlib import Path

def analyze_tables_structure(filepath):
    """Анализирует структуру таблиц."""
    doc = Document(filepath)
    
    print(f"\n{'='*60}")
    print(f"АНАЛИЗ: {filepath}")
    print(f"{'='*60}")
    print(f"Всего таблиц: {len(doc.tables)}\n")
    
    for t_idx, table in enumerate(doc.tables):
        print(f"--- Таблица {t_idx} ---")
        
        # Ширина таблицы
        tblPr = table._tbl.tblPr
        tbl_width = None
        tbl_type = None
        if tblPr is not None:
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tbl_width = tblW.get(qn('w:w'))
                tbl_type = tblW.get(qn('w:type'))
        print(f"  Ширина таблицы: {tbl_width} ({tbl_type})")
        print(f"  Строк: {len(table.rows)}")
        
        for r_idx, row in enumerate(table.rows):
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            
            height_val = None
            height_rule = None
            if trPr is not None:
                trH = trPr.find(qn('w:trHeight'))
                if trH is not None:
                    height_val = trH.get(qn('w:val'))
                    height_rule = trH.get(qn('w:hRule'))
            
            tcs = tr.findall(qn('w:tc'))
            cell_widths = []
            for tc in tcs:
                tcPr = tc.find(qn('w:tcPr'))
                w_val = None
                w_type = None
                if tcPr is not None:
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is not None:
                        w_val = tcW.get(qn('w:w'))
                        w_type = tcW.get(qn('w:type'))
                cell_widths.append(f"{w_val or 'AUTO'}({w_type or '-'})")
            
            print(f"    Строка {r_idx}: h={height_val}({height_rule}), ячеек={len(tcs)}")
            if r_idx < 3 or r_idx == len(table.rows) - 1:
                print(f"       Ширины: {cell_widths[:8]}{'...' if len(cell_widths) > 8 else ''}")
            elif r_idx == 3:
                print(f"    ...")

if __name__ == "__main__":
    if len(sys.argv) > 1:
        files = sys.argv[1:]
    else:
        files = [
            "contractor_report/болванки (шаблоны не вырезать только копировать)/Ежедневный отчет Шаблон.docx",
            "ЮНС_merged.docx"
        ]
    
    for f in files:
        p = Path(f)
        if p.exists():
            analyze_tables_structure(p)
        else:
            print(f"❌ Не найден: {p}")
