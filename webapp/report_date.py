"""
Единая логика даты: загруженные отчёты и болванки (текст + имя файла).
"""

import os
import re
from datetime import datetime, timedelta
from typing import Literal

from docx import Document

from docx_processing import _DATE_RE, replace_date_in_report_line as _replace_date_in_report_line_orig

_FILENAME_STEM_DATE_RE = re.compile(
    r"(\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4})\s*г\.?$",
    re.IGNORECASE,
)

_ZA_TITLE_TAIL = re.compile(
    r"(?P<prefix>.*?за)\s*"
    r"(?P<date>\d{1,2}\s*[./\-]\s*\d{1,2}\s*[./\-]\s*\d{2,4})\s*г\.?",
    re.IGNORECASE | re.DOTALL,
)


def resolve_target_date(
    mode: Literal["today", "yesterday"] | None = None,
    iso_date: str | None = None,
    ddmmyyyy: str | None = None,
) -> str:
    if ddmmyyyy:
        return ddmmyyyy
    if iso_date:
        return datetime.strptime(iso_date, "%Y-%m-%d").strftime("%d.%m.%Y")
    if mode == "yesterday":
        return (datetime.now() - timedelta(days=1)).strftime("%d.%m.%Y")
    return datetime.now().strftime("%d.%m.%Y")


def replace_date_in_filename(filename: str, target_date: str) -> str:
    stem, ext = os.path.splitext(filename)
    new_stem, n = _FILENAME_STEM_DATE_RE.subn(target_date + "г", stem, count=1)
    if n:
        return new_stem + ext
    new_stem2, n2 = _DATE_RE.subn(target_date, stem, count=1)
    return (new_stem2 if n2 else stem) + ext


def _write_paragraph_text(para, new_text: str) -> None:
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            r = run._element
            r.getparent().remove(r)
    else:
        para.add_run(new_text)


def update_template_paragraph_date(doc: Document, target_date: str) -> bool:
    for para in doc.paragraphs:
        full_text = "".join(run.text for run in para.runs) if para.runs else para.text
        if "за" not in full_text.lower():
            continue
        stripped = full_text.strip()
        m = _ZA_TITLE_TAIL.search(stripped)
        if m:
            new_text = m.group("prefix") + target_date + "г"
        elif stripped.endswith("за"):
            new_text = stripped + target_date + "г"
        elif _DATE_RE.search(stripped):
            new_text = _DATE_RE.sub(target_date, stripped, count=1)
            if not new_text.rstrip().endswith("г"):
                new_text = new_text.rstrip().rstrip(".") + "г"
        else:
            continue
        _write_paragraph_text(para, new_text)
        return True
    return False


def replace_date_in_report_line(
    doc: Document,
    mode: Literal["today", "yesterday"] = "today",
    target_date: str | None = None,
) -> bool:
    if target_date is None:
        target_date = resolve_target_date(mode=mode)
    return _replace_date_in_report_line_with_date(doc, target_date)


def _replace_date_in_report_line_with_date(doc: Document, target_date: str) -> bool:
    if len(doc.tables) > 0:
        table = doc.tables[0]
        if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
            cell = table.rows[2].cells[1]
            for para in cell.paragraphs:
                full_text = "".join(run.text for run in para.runs)
                if _DATE_RE.search(full_text):
                    new_text = _DATE_RE.sub(target_date, full_text)
                    if para.runs:
                        para.runs[0].text = new_text
                        for run in para.runs[1:]:
                            r = run._element
                            r.getparent().remove(r)
                    return True
    return False


def apply_date_to_templates(folder: str, target_date: str) -> list[str]:
    log = []
    for filename in sorted(os.listdir(folder)):
        if not filename.lower().endswith((".docx", ".doc")):
            continue
        filepath = os.path.join(folder, filename)
        try:
            doc = Document(filepath)
            if not update_template_paragraph_date(doc, target_date):
                log.append(f"Пропущен (нет строки «…за»): {filename}")
                continue
            doc.save(filepath)
            new_name = replace_date_in_filename(filename, target_date)
            if new_name != filename:
                dest = os.path.join(folder, new_name)
                if os.path.exists(dest) and os.path.realpath(dest) != os.path.realpath(filepath):
                    log.append(f"OK: {filename} (текст); имя не менял — уже есть {new_name}")
                else:
                    os.rename(filepath, dest)
                    log.append(f"OK: {filename} → {new_name}")
            else:
                log.append(f"OK: {filename} (дата в тексте → {target_date})")
        except Exception as e:
            log.append(f"Ошибка: {filename} — {e}")
    return log


def apply_date_to_uploads(folder: str, target_date: str) -> list[str]:
    log = []
    for filename in sorted(os.listdir(folder)):
        if not filename.lower().endswith((".docx", ".doc")):
            continue
        filepath = os.path.join(folder, filename)
        try:
            doc = Document(filepath)
            if not replace_date_in_report_line(doc, target_date=target_date):
                log.append(f"Пропущен (дата в таблице не найдена): {filename}")
                continue
            doc.save(filepath)
            new_name = replace_date_in_filename(filename, target_date)
            if new_name != filename:
                dest = os.path.join(folder, new_name)
                if os.path.exists(dest) and os.path.realpath(dest) != os.path.realpath(filepath):
                    log.append(f"OK: {filename} (дата в документе; имя занято: {new_name})")
                else:
                    os.rename(filepath, dest)
                    log.append(f"OK: {filename} → {new_name}")
            else:
                log.append(f"OK: {filename} (дата → {target_date})")
        except Exception as e:
            log.append(f"Ошибка: {filename} — {e}")
    return log


def apply_report_date(
    upload_dir: str,
    templates_dir: str,
    *,
    iso_date: str | None = None,
    mode: Literal["today", "yesterday"] | None = None,
) -> dict:
    target_date = resolve_target_date(mode=mode, iso_date=iso_date)
    iso_out = datetime.strptime(target_date, "%d.%m.%Y").strftime("%Y-%m-%d")
    return {
        "date": target_date,
        "iso_date": iso_out,
        "reports": apply_date_to_uploads(upload_dir, target_date),
        "templates": apply_date_to_templates(templates_dir, target_date),
    }


def rename_templates_compat(folder: str, mode: Literal["today", "yesterday"]) -> list[str]:
    return apply_date_to_templates(folder, resolve_target_date(mode=mode))
