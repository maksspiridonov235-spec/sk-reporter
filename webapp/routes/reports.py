import os
import shutil
from pathlib import Path

from docx import Document
from fastapi import APIRouter, File, HTTPException, UploadFile

from apply_template_layout import diagnose_document, hardcoded_layout
from config import UPLOAD_DIR

router = APIRouter()


@router.post("/upload/reports")
async def upload_reports(files: list[UploadFile] = File(...)):
    saved = []
    for f in files:
        if not f.filename:
            continue
        dest = UPLOAD_DIR / f.filename
        with open(dest, "wb") as out:
            shutil.copyfileobj(f.file, out)
        saved.append(f.filename)
    return {"uploaded": saved, "count": len(saved)}


@router.get("/files/reports")
async def list_reports():
    files = [f.name for f in UPLOAD_DIR.iterdir() if f.suffix.lower() in (".docx", ".doc")]
    return {"files": sorted(files)}


@router.get("/diagnose/reports", tags=["dev"], include_in_schema=False)
async def diagnose_reports():
    """DEV ONLY: диагностика сетки загруженных отчётов. В UI нет кнопки; только для отладки."""
    layout = hardcoded_layout()
    out = []
    for f in sorted(UPLOAD_DIR.iterdir()):
        if f.suffix.lower() not in (".docx", ".doc"):
            continue
        try:
            doc = Document(os.fspath(f))
            warns = diagnose_document(doc, layout)
            out.append({
                "file": f.name,
                "tables": len(doc.tables),
                "rows": [len(t.rows) for t in doc.tables],
                "images": len(doc.inline_shapes),
                "issues": warns,
                "ok": not warns,
            })
        except Exception as e:
            out.append({"file": f.name, "ok": False, "issues": [str(e)]})
    return {"dev_only": True, "reports": out, "grid_cols": layout["grid_cols"]}


@router.delete("/clear/reports")
async def clear_reports():
    shutil.rmtree(UPLOAD_DIR)
    UPLOAD_DIR.mkdir()
    return {"ok": True}


@router.post("/switch-leader-ai/{leader}")
async def switch_leader_ai_endpoint(leader: str):
    from agent.leader_ai_agent import switch_leader_ai

    if leader not in ("aniskov", "mandzhiev"):
        raise HTTPException(status_code=400, detail="leader должен быть 'aniskov' или 'mandzhiev'")
    report_files = list(UPLOAD_DIR.glob("*.docx"))
    if not report_files:
        raise HTTPException(status_code=404, detail="Отчёты не загружены")
    ok, msg = switch_leader_ai([str(f) for f in report_files], leader)
    if not ok:
        raise HTTPException(status_code=500, detail=msg)
    return {"ok": True, "message": msg}
