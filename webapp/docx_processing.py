"""
Логика обработки .docx файлов.
Все функции — аналоги VBA-макросов из Word, переписанные на python-docx.
"""

import os
import re
import shutil
import zipfile
import xml.etree.ElementTree as ET
from copy import deepcopy
from datetime import datetime, timedelta
from typing import Literal

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from lxml import etree

import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent.parent))
from companies import COMPANIES

# Regex для поиска дат в разных форматах (порядок важен!)
# Сначала 4-digit год, потом 2-digit, чтобы не обрезать YYYY на YY
_DATE_RE = re.compile(
    r"\d{1,2}\s*[./\-]\s*\d{1,2}\s*[./\-]\s*\d{4}\s*г\.?|"  # DD.MM.YYYY с буквой г (с пробелами)
    r"\d{1,2}[./\-]\d{1,2}[./\-]\d{4}\s*г\.?|"               # DD.MM.YYYY с буквой г
    r"\d{1,2}\s*[./\-]\s*\d{1,2}\s*[./\-]\s*\d{4}|"          # DD.MM.YYYY (с пробелами)
    r"\d{1,2}[./\-]\d{1,2}[./\-]\d{4}|"                      # DD.MM.YYYY
    r"\d{1,2}\s*[./\-]\s*\d{1,2}\s*[./\-]\s*\d{2}\s*г\.?|"   # DD.MM.YY с буквой г (с пробелами)
    r"\d{1,2}[./\-]\d{1,2}[./\-]\d{2}\s*г\.?|"               # DD.MM.YY с буквой г
    r"\d{1,2}\s*[./\-]\s*\d{1,2}\s*[./\-]\s*\d{2}|"          # DD.MM.YY (с пробелами)
    r"\d{1,2}[./\-]\d{1,2}[./\-]\d{2}"                       # DD.MM.YY
)


def _find_template_date(filename: str) -> str | None:
    """Возвращает дату из имени файла вида 'XX.XX.XXXX'."""
    m = _DATE_RE.search(filename)
    return m.group() if m else None


# ── Макрос 1: HighlightSecondRow_No5991 ────────────────────────────────────

def highlight_second_row(doc: Document) -> int:
    """
    Для каждой таблицы документа: если строк >= 2,
    красит 2-ю строку в голубой (#BDD6EE), жирный, по центру.
    Работает с объединёнными ячейками через XML напрямую.
    """
    BLUE_HEX = "BDD6EE"
    processed = 0

    for tbl in doc.tables:
        if len(tbl.rows) < 2:
            continue

        # Берём реальные <w:tc> элементы второй строки из XML (без дублей от merge)
        row_el = tbl.rows[1]._tr
        seen_tc = set()
        for tc in row_el.findall(qn("w:tc")):
            if id(tc) in seen_tc:
                continue
            seen_tc.add(id(tc))

            # Пропускаем ячейки с vMerge (продолжение вертикального объединения)
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is not None:
                vMerge = tcPr.find(qn("w:vMerge"))
                if vMerge is not None and vMerge.get(qn("w:val")) != "restart":
                    continue

            # Фон
            if tcPr is None:
                tcPr = etree.SubElement(tc, qn("w:tcPr"))
            shd = tcPr.find(qn("w:shd"))
            if shd is None:
                shd = etree.SubElement(tcPr, qn("w:shd"))
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), BLUE_HEX)
            # Убираем тему чтобы наш цвет не перекрывался
            for attr in list(shd.attrib):
                if "theme" in attr.lower():
                    del shd.attrib[attr]

            # Вертикальное выравнивание
            vAlign = tcPr.find(qn("w:vAlign"))
            if vAlign is None:
                vAlign = etree.SubElement(tcPr, qn("w:vAlign"))
            vAlign.set(qn("w:val"), "center")

            # Текст: жирный, по центру
            for p in tc.findall(qn("w:p")):
                pPr = p.find(qn("w:pPr"))
                if pPr is None:
                    pPr = etree.SubElement(p, qn("w:pPr"))
                jc = pPr.find(qn("w:jc"))
                if jc is None:
                    jc = etree.SubElement(pPr, qn("w:jc"))
                jc.set(qn("w:val"), "center")
                for r in p.findall(qn("w:r")):
                    rPr = r.find(qn("w:rPr"))
                    if rPr is None:
                        rPr = etree.SubElement(r, qn("w:rPr"))
                    b = rPr.find(qn("w:b"))
                    if b is None:
                        etree.SubElement(rPr, qn("w:b"))

        processed += 1

    return processed


# ── Макрос 2: NewMacros (форматирование документа) ─────────────────────────

def format_document(doc: Document) -> None:
    """
    - Шрифт Times New Roman 10pt для всего текста
    - Отступы и интервалы = 0, одинарный межстрочный
    - Ширина всех таблиц = 18.33 см, выравнивание по центру
    - Все инлайн-картинки → 5.33 × 4 см
    """
    # 18.33 см в единицах dxa (twentieths of a point): 1 cm = 567 dxa
    TABLE_WIDTH_DXA = int(18.33 * 567)
    # 5.33 см и 4 см в EMU (English Metric Units): 1 cm = 360000 EMU
    IMG_W_EMU = int(5.33 * 360000)
    IMG_H_EMU = int(4.0 * 360000)

    def _format_paras(paragraphs):
        for para in paragraphs:
            pPr = para._p.get_or_add_pPr()
            # Интервалы
            spacing = pPr.find(qn("w:spacing"))
            if spacing is None:
                spacing = etree.SubElement(pPr, qn("w:spacing"))
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            spacing.set(qn("w:line"), "240")
            spacing.set(qn("w:lineRule"), "auto")
            # Шрифт и размер для каждого run
            for run in para.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)
            # Если параграф вообще без runs — задаём через rPr по умолчанию
            rPr = pPr.find(qn("w:rPr"))
            if rPr is None:
                rPr = etree.SubElement(pPr, qn("w:rPr"))
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = etree.SubElement(rPr, qn("w:rFonts"))
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            sz = rPr.find(qn("w:sz"))
            if sz is None:
                sz = etree.SubElement(rPr, qn("w:sz"))
            sz.set(qn("w:val"), "20")  # 10pt = 20 half-points
            szCs = rPr.find(qn("w:szCs"))
            if szCs is None:
                szCs = etree.SubElement(rPr, qn("w:szCs"))
            szCs.set(qn("w:val"), "20")

    # Параграфы вне таблиц
    _format_paras(doc.paragraphs)

    # Таблицы
    for tbl in doc.tables:
        tblEl = tbl._tbl
        tblPr = tblEl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = etree.SubElement(tblEl, qn("w:tblPr"))

        # Ширина таблицы
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = etree.SubElement(tblPr, qn("w:tblW"))
        tblW.set(qn("w:w"), str(TABLE_WIDTH_DXA))
        tblW.set(qn("w:type"), "dxa")

        # Выравнивание по центру
        jc = tblPr.find(qn("w:jc"))
        if jc is None:
            jc = etree.SubElement(tblPr, qn("w:jc"))
        jc.set(qn("w:val"), "center")

        # Запрет авторастяжки
        autofit = tblPr.find(qn("w:tblLayout"))
        if autofit is None:
            autofit = etree.SubElement(tblPr, qn("w:tblLayout"))
        autofit.set(qn("w:type"), "fixed")

        for row in tbl.rows:
            for cell in row.cells:
                _format_paras(cell.paragraphs)

    # Картинки: меняем размер прямо в XML (cx/cy в EMU)
    EMU_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
    PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    for shape in doc.inline_shapes:
        # Находим <a:ext> внутри <wp:extent>
        drawing = shape._inline
        extent = drawing.find(
            "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent"
        )
        if extent is not None:
            extent.set("cx", str(IMG_W_EMU))
            extent.set("cy", str(IMG_H_EMU))
        # Также патчим <a:ext> внутри spPr
        for ext in drawing.iter(f"{{{EMU_NS}}}ext"):
            ext.set("cx", str(IMG_W_EMU))
            ext.set("cy", str(IMG_H_EMU))


# ── Макросы 3 и 4: ReplaceDateInReportLine / ReplaceDateInReportLine2 ──────

def replace_date_in_report_line(doc: Document, mode: Literal["today", "yesterday"]) -> bool:
    """
    Заменяет дату в ячейке [2,1] таблицы на сегодняшнюю или вчерашнюю.
    Работает для всех типов отчётов (геодезия, обычные, несгруппированные).
    Возвращает True если замена выполнена.
    """
    target_date = (
        datetime.now().strftime("%d.%m.%Y")
        if mode == "today"
        else (datetime.now() - timedelta(days=1)).strftime("%d.%m.%Y")
    )

    # Заменяем дату в ячейке [2,1] первой таблицы
    if len(doc.tables) > 0:
        table = doc.tables[0]
        if len(table.rows) > 2 and len(table.rows[2].cells) > 1:
            cell = table.rows[2].cells[1]
            for para in cell.paragraphs:
                # Собираем весь текст параграфа (дата может быть разбита на много runs)
                full_text = "".join(run.text for run in para.runs)
                if _DATE_RE.search(full_text):
                    # Заменяем дату, очищаем все runs и пишем результат в первый run
                    new_text = _DATE_RE.sub(target_date, full_text)
                    if para.runs:
                        para.runs[0].text = new_text
                        # Удаляем остальные runs
                        for run in para.runs[1:]:
                            r = run._element
                            r.getparent().remove(r)
                    return True

    return False


# ── Объединение отчётов ─────────────────────────────────────────────────────

def _zip_replace(zip_path: str, inner_name: str, new_data: bytes) -> None:
    """Заменяет файл внутри ZIP-архива."""
    tmp = zip_path + ".tmp"
    with zipfile.ZipFile(zip_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == inner_name:
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
    os.replace(tmp, zip_path)


_EXT_TO_CONTENT_TYPE = {
    ".jpeg": "image/jpeg",
    ".jpg":  "image/jpeg",
    ".png":  "image/png",
    ".gif":  "image/gif",
    ".bmp":  "image/bmp",
    ".tiff": "image/tiff",
    ".wmf":  "image/x-wmf",
    ".emf":  "image/x-emf",
}


def _copy_media_from_docx(src_path: str, dst_path: str, remap: dict) -> None:
    """
    Копирует медиафайлы из src в dst, обновляет [Content_Types].xml в dst.
    Заполняет remap {старое_имя: новое_имя}.
    """
    with zipfile.ZipFile(src_path, "r") as src_zip:
        src_media = [n for n in src_zip.namelist() if n.startswith("word/media/")]
        if not src_media:
            return

        # Читаем текущий [Content_Types].xml из dst
        with zipfile.ZipFile(dst_path, "r") as dst_zip:
            existing = set(dst_zip.namelist())
            ct_xml = dst_zip.read("[Content_Types].xml").decode("utf-8")

        ct_root = ET.fromstring(ct_xml)
        ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
        ET.register_namespace("", ct_ns)

        # Собираем уже объявленные расширения
        declared_exts = {
            el.get("Extension", "").lower()
            for el in ct_root.findall(f"{{{ct_ns}}}Default")
        }

        new_media: list[tuple[str, bytes]] = []  # (новое_имя_в_zip, данные)

        for src_name in src_media:
            base = os.path.basename(src_name)
            stem, ext = os.path.splitext(base)
            candidate = f"word/media/{base}"
            counter = 1
            while candidate in existing:
                candidate = f"word/media/{stem}_{counter}{ext}"
                counter += 1
            data = src_zip.read(src_name)
            new_media.append((candidate, data))
            existing.add(candidate)
            remap[base] = os.path.basename(candidate)

            # Добавляем Default в [Content_Types].xml если расширение новое
            ext_lower = ext.lower().lstrip(".")
            if ext_lower not in declared_exts:
                content_type = _EXT_TO_CONTENT_TYPE.get(ext.lower(), "application/octet-stream")
                new_default = ET.SubElement(ct_root, f"{{{ct_ns}}}Default")
                new_default.set("Extension", ext_lower)
                new_default.set("ContentType", content_type)
                declared_exts.add(ext_lower)

        # Записываем медиафайлы и обновлённый [Content_Types].xml в dst
        new_ct_xml = ET.tostring(ct_root, encoding="unicode", xml_declaration=False)
        new_ct_xml = "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>\n" + new_ct_xml

        with zipfile.ZipFile(dst_path, "a") as dst_zip:
            for name, data in new_media:
                dst_zip.writestr(name, data)

        _zip_replace(dst_path, "[Content_Types].xml", new_ct_xml.encode("utf-8"))


def _fix_image_refs_in_element(element, remap: dict) -> None:
    """
    Обновляет все ссылки на картинки в XML-элементе согласно remap.
    Картинки в docx ссылаются через r:embed в <a:blip> и <v:imagedata>.
    Remap применяется к именам файлов в relationships, но мы патчим
    уже resolved-имена прямо в XML через rId → имя файла.
    Этот патч применяется на уровне имён файлов в rels.
    """
    pass  # Rels патчим отдельно в merge_reports


def merge_reports(template_path: str, report_paths: list[str], output_path: str) -> int:
    """
    Объединяет отчёты из report_paths в шаблон template_path,
    сохраняет результат в output_path.
    Корректно переносит картинки через ZIP чтобы они не были битыми.
    Возвращает количество вставленных файлов.
    """
    shutil.copy2(template_path, output_path)
    inserted = 0

    for path in sorted(report_paths):
        try:
            src_doc = Document(path)
        except Exception:
            continue

        # 1. Копируем медиафайлы из src в output, получаем карту переименований
        remap: dict[str, str] = {}
        _copy_media_from_docx(path, output_path, remap)

        # 2. Читаем relationships из источника чтобы построить rId → имя файла
        src_rels: dict[str, str] = {}
        with zipfile.ZipFile(path, "r") as zf:
            rels_name = "word/_rels/document.xml.rels"
            if rels_name in zf.namelist():
                rels_xml = zf.read(rels_name).decode("utf-8")
                root = ET.fromstring(rels_xml)
                ns = "http://schemas.openxmlformats.org/package/2006/relationships"
                for rel in root.findall(f"{{{ns}}}Relationship"):
                    rid = rel.get("Id", "")
                    target = rel.get("Target", "")
                    # target вида "../media/image1.png" или "media/image1.png"
                    basename = os.path.basename(target)
                    src_rels[rid] = basename

        # 3. Добавляем relationships в output для новых медиафайлов
        #    и строим карту старых rId → новых rId
        rid_remap: dict[str, str] = {}
        if remap:
            with zipfile.ZipFile(output_path, "r") as zf:
                rels_name = "word/_rels/document.xml.rels"
                ET.register_namespace("", "http://schemas.openxmlformats.org/package/2006/relationships")
                rels_xml = zf.read(rels_name).decode("utf-8")
                rels_root = ET.fromstring(rels_xml)

            ns = "http://schemas.openxmlformats.org/package/2006/relationships"
            img_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"

            # Найдём максимальный rId в output
            existing_ids = set()
            for rel in rels_root.findall(f"{{{ns}}}Relationship"):
                existing_ids.add(rel.get("Id", ""))
            rid_counter = 1
            while f"rId{rid_counter}" in existing_ids:
                rid_counter += 1

            # Для каждого rId из источника с картинкой — добавляем новый rel в output
            for old_rid, old_basename in src_rels.items():
                if old_basename not in remap:
                    continue
                new_basename = remap[old_basename]
                new_rid = f"rId{rid_counter}"
                while new_rid in existing_ids:
                    rid_counter += 1
                    new_rid = f"rId{rid_counter}"
                rid_remap[old_rid] = new_rid
                existing_ids.add(new_rid)
                rid_counter += 1
                new_rel = ET.SubElement(rels_root, f"{{{ns}}}Relationship")
                new_rel.set("Id", new_rid)
                new_rel.set("Type", img_type)
                new_rel.set("Target", f"media/{new_basename}")

            # Записываем обновлённый rels обратно в ZIP
            new_rels_xml = ET.tostring(rels_root, encoding="unicode", xml_declaration=False)
            new_rels_xml = '<?xml version=\'1.0\' encoding=\'UTF-8\' standalone=\'yes\'?>\n' + new_rels_xml
            _zip_replace(output_path, rels_name, new_rels_xml.encode("utf-8"))

        # 4. Копируем тело документа, патча rId в XML-элементах
        master = Document(output_path)

        # Собираем элементы тела, пропуская ведущие пустые параграфы с разрывами
        body_elements = [
            el for el in src_doc.element.body
            if (el.tag.split("}")[-1] if "}" in el.tag else el.tag) != "sectPr"
        ]

        # Отрезаем пустые page-break параграфы в начале
        def _is_page_break_para(el):
            tag = el.tag.split("}")[-1] if "}" in el.tag else el.tag
            if tag != "p":
                return False
            from lxml import etree
            xml = etree.tostring(el, encoding="unicode")
            has_break = 'w:br' in xml and 'page' in xml
            # Текст — всё что не в тегах
            text = "".join(el.itertext()).strip()
            return has_break and not text

        while body_elements and _is_page_break_para(body_elements[0]):
            body_elements.pop(0)

        if inserted > 0 and body_elements:
            master.add_page_break()

        for element in body_elements:
            el_copy = deepcopy(element)
            if rid_remap:
                _patch_rids(el_copy, rid_remap)
            master.element.body.append(el_copy)

        master.save(output_path)
        inserted += 1

    return inserted


def _patch_rids(element, rid_remap: dict[str, str]) -> None:
    """Рекурсивно заменяет атрибуты r:embed, r:id, r:link на новые rId."""
    REMAP_ATTRS = (
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed",
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id",
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link",
    )
    for attr in REMAP_ATTRS:
        val = element.get(attr)
        if val and val in rid_remap:
            element.set(attr, rid_remap[val])
    for child in element:
        _patch_rids(child, rid_remap)


# ── Переименование файлов ───────────────────────────────────────────────────

def rename_files(folder: str, mode: Literal["today", "yesterday"]) -> list[str]:
    """
    Переименовывает .docx/.doc файлы в папке:
    заменяет все даты вида DD.MM.YYYY в имени файла на сегодняшнюю или вчерашнюю.
    Возвращает список строк лога.
    """
    new_date = (
        datetime.now().strftime("%d.%m.%Y")
        if mode == "today"
        else (datetime.now() - timedelta(days=1)).strftime("%d.%m.%Y")
    )
    log = []
    for filename in os.listdir(folder):
        if not any(filename.lower().endswith(ext) for ext in (".docx", ".doc")):
            continue
        new_name = _DATE_RE.sub(new_date, filename)
        if new_name != filename:
            try:
                os.rename(
                    os.path.join(folder, filename),
                    os.path.join(folder, new_name),
                )
                log.append(f"Переименован: {filename} → {new_name}")
            except Exception as e:
                log.append(f"Ошибка: {filename} — {e}")
    return log


# ── Переименование результатов ──────────────────────────────────────────────

def rename_results(folder: str, mode: Literal["today", "yesterday"]) -> list[str]:
    """
    Переименовывает Евракор_merged.docx → Евракор_Ежедневный отчёт СК за DD.MM.YYYY.docx
    """
    new_date = (
        datetime.now().strftime("%d.%m.%Y")
        if mode == "today"
        else (datetime.now() - timedelta(days=1)).strftime("%d.%m.%Y")
    )
    log = []
    for filename in os.listdir(folder):
        if not filename.lower().endswith((".docx", ".doc")):
            continue
        if "_merged" not in filename:
            continue
        company = filename.replace("_merged.docx", "").replace("_merged.doc", "")
        ext = ".docx" if filename.lower().endswith(".docx") else ".doc"
        new_name = f"{company}_Ежедневный отчёт СК за {new_date}{ext}"
        filepath = os.path.join(folder, filename)
        try:
            # Заменяем дату внутри документа
            doc = Document(filepath)
            replace_date_in_report_line(doc, mode)
            doc.save(filepath)
            # Переименовываем файл
            os.rename(filepath, os.path.join(folder, new_name))
            log.append(f"Переименован: {filename} → {new_name}")
        except Exception as e:
            log.append(f"Ошибка: {filename} — {e}")
    return log


def rename_templates(folder: str, mode: Literal["today", "yesterday"]) -> list[str]:
    """
    Переименовывает шаблоны и меняет дату внутри них.
    НГП_Ежедневный отчёт СК за 02.08.2025г.docx → НГП_Ежедневный отчёт СК за DD.MM.YYYY.docx
    """
    new_date = (
        datetime.now().strftime("%d.%m.%Y")
        if mode == "today"
        else (datetime.now() - timedelta(days=1)).strftime("%d.%m.%Y")
    )
    log = []
    for filename in os.listdir(folder):
        if not filename.lower().endswith((".docx", ".doc")):
            continue
        filepath = os.path.join(folder, filename)
        try:
            doc = Document(filepath)
            old_date = _find_template_date(filename)

            if not old_date:
                log.append(f"Пропущен (дата не найдена): {filename}")
                continue

            # Меняем дату в параграфах
            for para in doc.paragraphs:
                if old_date in para.text:
                    para.text = para.text.replace(old_date, new_date)

            # Меняем дату в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if old_date in para.text:
                                para.text = para.text.replace(old_date, new_date)

            doc.save(filepath)

            # Переименовываем файл
            new_name = filename.replace(old_date, new_date)
            new_filepath = os.path.join(folder, new_name)
            if filepath != new_filepath:
                os.rename(filepath, new_filepath)

            log.append(f"Обновлён: {filename} → {new_name}")
        except Exception as e:
            log.append(f"Ошибка: {filename} — {e}")
    return log


# ── Применение макроса к файлу ──────────────────────────────────────────────

def apply_macro_to_file(filepath: str, macro_name: str) -> tuple[bool, str]:
    """
    Применяет один из макросов к файлу, сохраняет его.
    Возвращает (успех, сообщение).
    """
    try:
        doc = Document(filepath)
    except Exception as e:
        return False, str(e)

    if macro_name == "HighlightSecondRow_No5991":
        n = highlight_second_row(doc)
        msg = f"Обработано таблиц: {n}"
    elif macro_name == "NewMacros":
        format_document(doc)
        msg = "Форматирование применено"
    elif macro_name == "ReplaceDateInReportLine":
        ok = replace_date_in_report_line(doc, "today")
        msg = "Дата заменена на сегодняшнюю" if ok else "Строка с датой не найдена"
    elif macro_name == "ReplaceDateInReportLine2":
        ok = replace_date_in_report_line(doc, "yesterday")
        msg = "Дата заменена на вчерашнюю" if ok else "Строка с датой не найдена"
    else:
        return False, f"Неизвестный макрос: {macro_name}"

    doc.save(filepath)
    return True, msg
