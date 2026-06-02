"""
Применяет сетку столбцов из эталонного шаблона ко всем таблицам документа.
"""

import os
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

BOLVANKI_DIR = (
    Path(__file__).parent
    / "contractor_report"
    / "болванки (шаблоны не вырезать только копировать)"
)

# 6 колонок — исходный эталон (Пряхин и большинство отчётов)
DEFAULT_GRID_COLS_6 = ["2041", "1757", "1787", "1898", "1701", "1646"]

# 7 колонок — целевая сетка Громова для всех таблиц после prepare/merge
# Первые 5 = как эталон, 6+7 = блок договора (1291+355 = 1646)
_GROMOV_CONTRACT_LABEL_DXA = "1291"
_GROMOV_CONTRACT_VALUE_DXA = "355"
DEFAULT_GRID_COLS = list(DEFAULT_GRID_COLS_6[:5]) + [
    _GROMOV_CONTRACT_LABEL_DXA,
    _GROMOV_CONTRACT_VALUE_DXA,
]

# Эталон 6-кол. — образец раскладки строк (источник для 6→7)
ETALON_REPORT_PATH = (
    Path(__file__).parent
    / "Ежедневный отчет (ЮНС) от 26.04.2026 г. (БКНС-4) Пряхин И.Н..docx"
)

# gridSpan 7→6 (справочно); для prepare используем обратное 6→7
_SPAN_PATTERN_7_TO_6: dict[tuple[int, ...], tuple[int, ...]] = {
    (4, 1, 1, 1): (3, 1, 1, 1),
    (1, 2, 1, 3): (1, 1, 1, 3),
    (5, 2): (4, 2),
    (1, 4, 1, 1): (1, 3, 1, 1),
    (7,): (6,),
    (3, 2, 2): (2, 2, 2),
    (1, 6): (1, 5),
    (5, 1, 1): (4, 1, 1),
    (3, 4): (2, 4),
    (1, 2, 1, 1, 1, 1): (1, 1, 1, 1, 1, 1),
    (1, 3, 2, 1): (1, 2, 2, 1),
    (1, 4, 2): (1, 3, 2),
    (3, 1, 1, 1, 1): (2, 1, 1, 1, 1),
    (2, 3, 1, 1): (2, 2, 1, 1),
}
_SPAN_PATTERN_6_TO_7: dict[tuple[int, ...], tuple[int, ...]] = {
    v: k for k, v in _SPAN_PATTERN_7_TO_6.items()
}

_etalon_spans_by_signature: dict[str, tuple[int, ...]] | None = None
ROW_HEIGHT = "340"
ROW_HEIGHT_RULE = "atLeast"
MIN_ROW_HEIGHT_CM = 0.6
MIN_ROWS_FOR_MAIN_TABLE = 8
TARGET_TABLE_WIDTH = sum(int(w) for w in DEFAULT_GRID_COLS)


def hardcoded_layout() -> dict:
    return {
        "template": "hardcoded",
        "grid_cols": list(DEFAULT_GRID_COLS),
        "grid_cols_6": list(DEFAULT_GRID_COLS_6),
        "tblGrid": None,
    }


def resolve_layout_template(templates_dir: Path | None = None) -> Path:
    templates_dir = Path(templates_dir or BOLVANKI_DIR)
    if not templates_dir.is_dir():
        raise FileNotFoundError(f"Папка болванок не найдена: {templates_dir}")

    for name in ("Ежедневный отчет Шаблон.docx", "Ежедневный отчёт Шаблон.docx"):
        p = templates_dir / name
        if p.is_file():
            return p

    for pattern in ("*Шаблон*.docx", "*шаблон*.docx"):
        found = sorted(templates_dir.glob(pattern))
        if found:
            return found[0]

    for pattern in ("ЮНС_*.docx", "Евракор_*.docx"):
        found = sorted(templates_dir.glob(pattern))
        if found:
            return found[0]

    any_docx = sorted(templates_dir.glob("*.docx"))
    if any_docx:
        return any_docx[0]

    raise FileNotFoundError(f"В {templates_dir} нет ни одного .docx для сетки таблицы")


def _grid_cols_from_tbl(tbl) -> list[str]:
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is None:
        return list(DEFAULT_GRID_COLS)
    cols = []
    for col in tblGrid.findall(qn("w:gridCol")):
        w = col.get(qn("w:w"))
        if w:
            cols.append(w)
    return cols or list(DEFAULT_GRID_COLS)


def read_template_layout(template_path: Path) -> dict:
    doc = Document(os.fspath(template_path))
    if not doc.tables:
        raise ValueError(f"В шаблоне нет таблиц: {template_path}")
    tbl = doc.tables[0]._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    grid_cols = _grid_cols_from_tbl(tbl)
    return {
        "template": str(template_path),
        "tblGrid": deepcopy(tblGrid) if tblGrid is not None else None,
        "grid_cols": grid_cols,
    }


def _build_tblGrid(grid_cols: list[str]) -> etree._Element:
    tblGrid = etree.Element(qn("w:tblGrid"))
    for w in grid_cols:
        col = etree.SubElement(tblGrid, qn("w:gridCol"))
        col.set(qn("w:w"), w)
    return tblGrid


def _grid_cumsum(grid_cols: list[str]) -> list[int]:
    cs = [0]
    for w in grid_cols:
        cs.append(cs[-1] + int(w))
    return cs


def _row_start_col(tr) -> int:
    tr_pr = tr.find(qn("w:trPr"))
    if tr_pr is None:
        return 0
    gb = tr_pr.find(qn("w:gridBefore"))
    if gb is None:
        return 0
    return int(gb.get(qn("w:val"), 0))


def _row_occupied_cols(tr, ncol: int) -> int:
    col_idx = _row_start_col(tr)
    for tc in tr.findall(qn("w:tc")):
        if col_idx >= ncol:
            break
        tc_pr = tc.find(qn("w:tcPr"))
        span = 1
        if tc_pr is not None:
            gs = tc_pr.find(qn("w:gridSpan"))
            if gs is not None:
                span = max(1, int(gs.get(qn("w:val"), 1)))
            vm = tc_pr.find(qn("w:vMerge"))
            if vm is not None and vm.get(qn("w:val")) != "restart":
                col_idx += span
                continue
        col_idx += span
    return col_idx


def _normalize_row_tr(tr) -> None:
    """Убирает gridBefore — частый артефакт, ломает подсчёт колонок."""
    tr_pr = tr.find(qn("w:trPr"))
    if tr_pr is None:
        return
    gb = tr_pr.find(qn("w:gridBefore"))
    if gb is not None:
        tr_pr.remove(gb)


def _set_tc_width(tc, width_dxa: str) -> None:
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tc_pr)
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = etree.SubElement(tc_pr, qn("w:tcW"))
    tc_w.set(qn("w:w"), width_dxa)
    tc_w.set(qn("w:type"), "dxa")


def _set_grid_span(tc, span: int) -> None:
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tc_pr)
    gs = tc_pr.find(qn("w:gridSpan"))
    if span <= 1:
        if gs is not None:
            tc_pr.remove(gs)
        return
    if gs is None:
        gs = etree.SubElement(tc_pr, qn("w:gridSpan"))
    gs.set(qn("w:val"), str(span))


def _tc_is_vmerge_continue(tc) -> bool:
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return False
    vm = tc_pr.find(qn("w:vMerge"))
    return vm is not None and vm.get(qn("w:val")) != "restart"


def _get_grid_span(tc) -> int:
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        return 1
    gs = tc_pr.find(qn("w:gridSpan"))
    if gs is None:
        return 1
    return max(1, int(gs.get(qn("w:val"), 1)))



def _col6_index_to_col7(col6: int) -> int:
    """Обратное к 7→6: между 6c1 и 6c2 вставляется доп. gridCol в 7-кол."""
    if col6 <= 0:
        return 0
    if col6 == 1:
        return 1
    return col6 + 1


def _remap_span_6_to_7(col6_start: int, span6: int) -> int:
    end6 = col6_start + span6 - 1
    c7_start = _col6_index_to_col7(col6_start)
    c7_end = _col6_index_to_col7(end6)
    return c7_end - c7_start + 1


def _row_span_pattern(tr) -> tuple[int, ...]:
    spans: list[int] = []
    for tc in tr.findall(qn("w:tc")):
        if _tc_is_vmerge_continue(tc):
            continue
        spans.append(_get_grid_span(tc))
    return tuple(spans)


def _row_span_pattern_7(tr) -> tuple[int, ...] | None:
    spans = _row_span_pattern(tr)
    if not spans or sum(spans) != 7:
        return None
    return spans


def _row_signature(tr) -> str:
    parts: list[str] = []
    for tc in tr.findall(qn("w:tc")):
        if _tc_is_vmerge_continue(tc):
            continue
        text = "".join(t.text for t in tc.iter(qn("w:t")) if t.text).strip()
        parts.append(text[:40].lower())
    return "||".join(parts)


def _load_etalon_span_signatures() -> dict[str, tuple[int, ...]]:
    global _etalon_spans_by_signature
    if _etalon_spans_by_signature is not None:
        return _etalon_spans_by_signature

    out: dict[str, tuple[int, ...]] = {}
    if ETALON_REPORT_PATH.is_file():
        doc = Document(os.fspath(ETALON_REPORT_PATH))
        if doc.tables:
            for tr in doc.tables[0]._tbl.findall(qn("w:tr")):
                spans = _row_span_pattern(tr)
                if not spans or sum(spans) != 6:
                    continue
                sig = _row_signature(tr)
                out[sig] = spans
                if parts := sig.split("||"):
                    out[parts[0]] = spans
                    if len(parts) > 1:
                        out[parts[0] + "||" + parts[1]] = spans

    _etalon_spans_by_signature = out
    return out


def _row_span_pattern_6(tr) -> tuple[int, ...] | None:
    spans = _row_span_pattern(tr)
    if not spans or sum(spans) != 6:
        return None
    return spans


def _target_spans_7_for_6col_row(tr, pattern6: tuple[int, ...], n_cells: int) -> tuple[int, ...] | None:
    if pattern6 in _SPAN_PATTERN_6_TO_7:
        target = _SPAN_PATTERN_6_TO_7[pattern6]
        if len(target) == n_cells:
            return target

    parts = _row_signature(tr).split("||")
    head = parts[0] if parts else ""

    if "статус строительного" in head:
        return (3, 1, 1, 1, 1) if n_cells == 5 else None
    if "описание действий" in head or (head and head[0].isdigit() and "." in head[:4]):
        return (5, 1, 1) if n_cells == 3 else None
    if "погодная характеристика" in head:
        return (1, 4, 2) if n_cells == 3 else None
    if head.startswith("+") and "°" in head:
        return (1, 4, 2) if n_cells == 3 else None
    if "результат строительного" in head:
        return (7,) if n_cells == 1 else None
    if "данные заказчика" in head:
        return (5, 2) if n_cells == 2 else None
    if head in ("тел.:", "тел:"):
        return (1, 3, 2, 1) if n_cells == 4 else None
    if not head.strip() and pattern6 == (2, 2, 1, 1):
        return (2, 3, 1, 1)

    return None


def _normalize_6col_table_to_7col_grid(table) -> bool:
    """6-кол. отчёт → раскладка строк как у Громова (7 gridCol)."""
    changed = False
    for row in table.rows:
        tr = row._tr
        _normalize_row_tr(tr)
        if _row_occupied_cols(tr, 8) != 6:
            continue

        tcs = [tc for tc in tr.findall(qn("w:tc")) if not _tc_is_vmerge_continue(tc)]
        if not tcs:
            continue

        pattern6 = _row_span_pattern_6(tr)
        target = None
        if pattern6 is not None:
            target = _target_spans_7_for_6col_row(tr, pattern6, len(tcs))
            if target is not None and len(target) != len(tcs):
                target = None

        if len(tcs) == 2 and sum(_get_grid_span(tc) for tc in tcs) == 4:
            # подписи с vMerge: две видимые ячейки span2+span2 → 4+3
            _set_grid_span(tcs[0], 4)
            _set_grid_span(tcs[1], 3)
        else:
            col6 = 0
            for i, tc in enumerate(tcs):
                span6 = _get_grid_span(tc)
                if target is not None:
                    span7 = target[i]
                else:
                    span7 = _remap_span_6_to_7(col6, span6)
                _set_grid_span(tc, span7)
                col6 += span6
        changed = True

    return changed


def _max_row_occupancy(tbl) -> int:
    max_occ = 0
    for tr in tbl.findall(qn("w:tr")):
        max_occ = max(max_occ, _row_occupied_cols(tr, 16))
    return max_occ


def _grid_cols_for_table(tbl, layout_cols: list[str] | None = None) -> list[str]:
    """Целевая сетка — 7 колонок (Громов) для всех отчётных таблиц."""
    return list(DEFAULT_GRID_COLS)


def _table_needs_6_to_7_normalize(tbl) -> bool:
    """Есть 6-колоночные строки — разворачиваем в 7 (сетка Громова)."""
    if _max_row_occupancy(tbl) >= 7 and len(_grid_cols_from_tbl(tbl)) >= 7:
        return False
    return _max_row_occupancy(tbl) <= 6 or len(_grid_cols_from_tbl(tbl)) <= 6


def diagnose_table(tbl, grid_cols: list[str]) -> list[str]:
    ncol = len(grid_cols)
    issues: list[str] = []
    rows = tbl.findall(qn("w:tr"))
    tbl_grid = tbl.find(qn("w:tblGrid"))
    file_ncol = len(tbl_grid.findall(qn("w:gridCol"))) if tbl_grid is not None else 0
    if file_ncol and file_ncol != ncol:
        issues.append(f"в файле {file_ncol} колонок сетки, ожид. {ncol}")
    bad = []
    for ri, tr in enumerate(rows):
        occ = _row_occupied_cols(tr, ncol)
        if occ != ncol:
            bad.append(f"{ri + 1}({occ})")
    if bad:
        preview = ", ".join(bad[:6])
        more = f" +{len(bad) - 6}" if len(bad) > 6 else ""
        issues.append(f"битые строки: {preview}{more}")
    return issues


def diagnose_document(doc, layout: dict | None = None) -> list[str]:
    standard = list(DEFAULT_GRID_COLS)
    out: list[str] = []
    if not doc.tables:
        out.append("нет таблиц")
        return out
    for i, table in enumerate(doc.tables):
        tbl = table._tbl
        nrows = len(table.rows)
        expected = _grid_cols_for_table(tbl)
        issues = diagnose_table(tbl, expected)
        if issues:
            out.append(f"табл.{i + 1} ({nrows} стр.): " + "; ".join(issues))
    return out


def _main_table_indices(doc) -> list[int]:
    """Сетку на все таблицы-отчёты (>= MIN_ROWS), в т.ч. в merged.docx."""
    if not doc.tables:
        return []
    scored = [(i, len(t.rows)) for i, t in enumerate(doc.tables)]
    qualifying = [i for i, n in scored if n >= MIN_ROWS_FOR_MAIN_TABLE]
    if qualifying:
        return qualifying
    return [i for i, n in scored if n >= 3] or [0]


def apply_layout(doc, layout: dict | None = None, only_main_table: bool = True) -> list[str]:
    """Возвращает предупреждения по документу."""
    target_cols = list(DEFAULT_GRID_COLS)
    indices = _main_table_indices(doc) if only_main_table else list(range(len(doc.tables)))

    for ti in indices:
        table = doc.tables[ti]
        tbl = table._tbl
        if _table_needs_6_to_7_normalize(tbl):
            _normalize_6col_table_to_7col_grid(table)

        grid_cols = _grid_cols_for_table(tbl)
        cumsum = _grid_cumsum(grid_cols)
        table_width = str(cumsum[-1])

        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = etree.SubElement(tbl, qn("w:tblPr"))
            tbl.insert(0, tbl_pr)

        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is None:
            tbl_w = etree.SubElement(tbl_pr, qn("w:tblW"))
        tbl_w.set(qn("w:w"), table_width)
        tbl_w.set(qn("w:type"), "dxa")

        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = etree.SubElement(tbl_pr, qn("w:tblLayout"))
        tbl_layout.set(qn("w:type"), "fixed")

        jc = tbl_pr.find(qn("w:jc"))
        if jc is None:
            jc = etree.SubElement(tbl_pr, qn("w:jc"))
        jc.set(qn("w:val"), "center")

        old_grid = tbl.find(qn("w:tblGrid"))
        new_grid = _build_tblGrid(grid_cols)
        if old_grid is not None:
            tbl.replace(old_grid, new_grid)
        else:
            tbl_pr.addnext(new_grid)

        for row in table.rows:
            tr = row._tr
            _normalize_row_tr(tr)
            col_idx = 0

            for tc in tr.findall(qn("w:tc")):
                if col_idx >= len(grid_cols):
                    break

                tc_pr = tc.find(qn("w:tcPr"))
                span = 1
                if tc_pr is not None:
                    gs = tc_pr.find(qn("w:gridSpan"))
                    if gs is not None:
                        span = max(1, int(gs.get(qn("w:val"), 1)))
                span = min(span, len(grid_cols) - col_idx)

                cell_w = str(cumsum[col_idx + span] - cumsum[col_idx])
                # Важно: ширину задаём и vMerge-continuation, иначе остаётся старый tcW
                _set_tc_width(tc, cell_w)
                col_idx += span

    warnings = diagnose_document(doc, layout)
    return warnings


def main():
    if len(sys.argv) < 2:
        print("Использование: python3 apply_template_layout.py <document.docx> [эталон.docx]")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    template_path = Path(sys.argv[2]) if len(sys.argv) > 2 else resolve_layout_template()

    layout = read_template_layout(template_path)
    doc = Document(os.fspath(input_path))
    print(f"Эталон: {template_path.name}, колонки: {layout['grid_cols']}")
    warns = apply_layout(doc, layout)
    for w in warns:
        print("WARN:", w)

    output_path = input_path.parent / f"{input_path.stem}_layout.docx"
    doc.save(os.fspath(output_path))
    print(f"Сохранён: {output_path}")


if __name__ == "__main__":
    main()
